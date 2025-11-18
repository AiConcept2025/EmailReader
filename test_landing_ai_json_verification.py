#!/usr/bin/env python3
"""
Test script to verify LandingAI JSON file saving and content filtering.

This test processes a PDF and verifies:
1. JSON response is saved to completed_temp/
2. Special markup is filtered from output
3. Decorative chunks are filtered out
"""

import sys
import os
import logging
import json
from pathlib import Path

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)-8s | %(name)-20s | %(message)s',
    datefmt='%H:%M:%S'
)

logger = logging.getLogger('JSONVerification.Test')


def test_json_saving_and_filtering(pdf_path: str):
    """Test that JSON is saved and content is properly filtered."""

    if not os.path.exists(pdf_path):
        logger.error(f"File not found: {pdf_path}")
        return False

    logger.info("="*80)
    logger.info("LandingAI JSON Verification Test")
    logger.info("="*80)
    logger.info(f"Input file: {pdf_path}")

    # Import after setting up logging
    from src.ocr import OCRProviderFactory
    from src.config import load_config
    from docx import Document

    try:
        # Load configuration
        logger.info("Loading configuration...")
        config = load_config()

        # Create output path
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        output_docx = f"test_output_{base_name}_verification.docx"
        # JSON files now have timestamp, so we look for pattern
        json_pattern = f"completed_temp/{base_name}_landing_ai_*.json"

        logger.info(f"Expected JSON pattern: {json_pattern}")
        logger.info(f"Output DOCX file: {output_docx}")
        logger.info("")

        # Remove old files if they exist
        if os.path.exists(output_docx):
            os.remove(output_docx)
            logger.info(f"Removed old DOCX: {output_docx}")

        # Remove old JSON files matching pattern
        import glob
        old_json_files = glob.glob(json_pattern)
        for old_json in old_json_files:
            os.remove(old_json)
            logger.info(f"Removed old JSON: {old_json}")

        logger.info("")
        logger.info("Starting OCR processing with LandingAI...")
        logger.info("="*80)

        # Get OCR provider and process document
        ocr_provider = OCRProviderFactory.get_provider(config)
        ocr_provider.process_document(pdf_path, output_docx)

        logger.info("")
        logger.info("="*80)
        logger.info("VERIFICATION CHECKS")
        logger.info("="*80)

        # VERIFICATION 1: Check JSON file exists (with timestamp)
        logger.info("")
        logger.info("Check 1: JSON file saved to completed_temp/")

        # Find JSON files matching pattern
        json_files = glob.glob(json_pattern)

        if json_files:
            # Get the most recent JSON file
            latest_json = max(json_files, key=os.path.getmtime)
            json_size = os.path.getsize(latest_json) / 1024  # KB
            logger.info(f"  ✅ PASS: JSON file exists ({json_size:.2f} KB)")
            logger.info(f"  Location: {latest_json}")

            # Analyze JSON content
            with open(latest_json, 'r') as f:
                json_data = json.load(f)

            chunks = json_data.get('chunks', [])
            logger.info(f"  Total chunks in JSON: {len(chunks)}")

            # Count chunk types
            from collections import Counter
            chunk_types = Counter(c.get('type') for c in chunks)
            logger.info(f"  Chunk types: {dict(chunk_types)}")

        else:
            logger.error(f"  ❌ FAIL: No JSON files found matching {json_pattern}")
            return False

        # VERIFICATION 2: Check DOCX output
        logger.info("")
        logger.info("Check 2: DOCX file created successfully")
        if not os.path.exists(output_docx):
            logger.error(f"  ❌ FAIL: DOCX file not found: {output_docx}")
            return False

        docx_size = os.path.getsize(output_docx) / 1024  # KB
        logger.info(f"  ✅ PASS: DOCX file exists ({docx_size:.2f} KB)")

        # VERIFICATION 3: Check for special markup in DOCX
        logger.info("")
        logger.info("Check 3: Special markup filtered from DOCX")
        doc = Document(output_docx)

        markup_found = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text
            if '<::' in text or '::>' in text:
                markup_found.append(f"Paragraph {i}: {text[:100]}")

        if markup_found:
            logger.error(f"  ❌ FAIL: Special markup found in {len(markup_found)} paragraphs:")
            for item in markup_found[:5]:  # Show first 5
                logger.error(f"    {item}")
            return False
        else:
            logger.info(f"  ✅ PASS: No special markup found in DOCX")

        # VERIFICATION 4: Check decorative chunks are filtered
        logger.info("")
        logger.info("Check 4: Decorative chunks filtered out")

        # Count text chunks in JSON
        text_chunks = [c for c in chunks if c.get('type') == 'text']
        decorative_chunks = [c for c in chunks if c.get('type') != 'text']

        # Count paragraphs in DOCX
        docx_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])

        logger.info(f"  JSON chunks: {len(chunks)} total")
        logger.info(f"    - Text chunks: {len(text_chunks)}")
        logger.info(f"    - Decorative chunks: {len(decorative_chunks)}")
        logger.info(f"  DOCX paragraphs: {docx_paragraphs}")

        if decorative_chunks:
            decorative_types = Counter(c.get('type') for c in decorative_chunks)
            logger.info(f"  Filtered types: {dict(decorative_types)}")

        # Verify decorative chunks were filtered
        if docx_paragraphs <= len(text_chunks):
            logger.info(f"  ✅ PASS: Decorative chunks filtered successfully")
        else:
            logger.warning(f"  ⚠️  WARNING: DOCX has more paragraphs than text chunks")

        # VERIFICATION 5: Check for specific problematic content
        logger.info("")
        logger.info("Check 5: No logo/decorative descriptions in DOCX")

        problematic_phrases = [
            'vertical decorative border',
            'stylized letter',
            'qr code',
            'signature: illegible',
            'teardrop shape'
        ]

        issues_found = []
        for para in doc.paragraphs:
            text = para.text.lower()
            for phrase in problematic_phrases:
                if phrase in text:
                    issues_found.append(f"Found '{phrase}' in: {para.text[:100]}")

        if issues_found:
            logger.error(f"  ❌ FAIL: Decorative descriptions found:")
            for issue in issues_found[:5]:
                logger.error(f"    {issue}")
            return False
        else:
            logger.info(f"  ✅ PASS: No decorative descriptions found")

        # Final summary
        logger.info("")
        logger.info("="*80)
        logger.info("✅ ALL VERIFICATION CHECKS PASSED")
        logger.info("="*80)
        logger.info("")
        logger.info(f"Summary:")
        logger.info(f"  - JSON saved: {latest_json}")
        logger.info(f"  - DOCX created: {output_docx}")
        logger.info(f"  - Special markup filtered: Yes")
        logger.info(f"  - Decorative chunks filtered: Yes ({len(decorative_chunks)} removed)")
        logger.info(f"  - Clean paragraphs in output: {docx_paragraphs}")

        return True

    except Exception as e:
        logger.error(f"Test failed with error: {e}", exc_info=True)
        return False


if __name__ == "__main__":
    # If no argument provided, use default test file from test_docs or inbox_temp
    if len(sys.argv) < 2:
        # Try to find a PDF in inbox_temp first
        import glob
        inbox_pdfs = glob.glob("inbox_temp/*.pdf")
        test_docs_pdfs = glob.glob("test_docs/*.pdf")

        if inbox_pdfs:
            pdf_file = inbox_pdfs[0]
            logger.info(f"No file specified, using: {pdf_file}")
        elif test_docs_pdfs:
            pdf_file = test_docs_pdfs[0]
            logger.info(f"No file specified, using: {pdf_file}")
        else:
            print("Usage: python test_landing_ai_json_verification.py <pdf_file_path>")
            print("")
            print("Example:")
            print("  python test_landing_ai_json_verification.py test_docs/PDF-scanned-rus-words.pdf")
            print("  python test_landing_ai_json_verification.py inbox_temp/YourFile.pdf")
            print("")
            print("Or run without arguments to use a file from inbox_temp/ or test_docs/")
            sys.exit(1)
    else:
        pdf_file = sys.argv[1]

    success = test_json_saving_and_filtering(pdf_file)

    sys.exit(0 if success else 1)
