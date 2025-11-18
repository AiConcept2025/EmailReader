#!/usr/bin/env python3
"""
Iterative Test Script for EmailReader Document Formatting Validation

This script validates that document formatting is preserved through OCR and translation:
1. Number of pages matches original PDF
2. No bold or italic fonts
3. All text is Times New Roman 12pt
4. All formatting preserved from OCR through translation

The test runs iteratively up to 10 times, fixing issues as they are found.
"""

import os
import sys
import subprocess
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from docx import Document
from docx.shared import Pt

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Project paths
PROJECT_ROOT = Path(__file__).parent
GOOGLE_TRANSLATOR_ROOT = PROJECT_ROOT.parent / "GoogleTranslator"
TEST_DOCS_DIR = PROJECT_ROOT / "test_docs"
INBOX_TEMP_DIR = PROJECT_ROOT / "inbox_temp"


class DocumentAnalyzer:
    """Analyzes DOCX files for formatting compliance."""

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.issues: List[str] = []

    def count_pages(self) -> int:
        """
        Count pages in DOCX document.
        Note: python-docx doesn't have direct page count,
        so we count sections and page breaks as a proxy.
        """
        page_count = 1  # Start with 1 page

        # Count explicit page breaks in regular paragraphs
        for paragraph in self.doc.paragraphs:
            if paragraph.paragraph_format.page_break_before:
                page_count += 1

        # Check for page breaks in runs
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                # Check for page break character
                if '\f' in run.text or '\x0c' in run.text:
                    page_count += 1

        # Check for page breaks in table cells
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.paragraph_format.page_break_before:
                            page_count += 1

        return page_count

    def analyze_formatting(self) -> Dict:
        """
        Analyze all formatting in the document.

        Returns:
            Dictionary with analysis results
        """
        results = {
            'page_count': self.count_pages(),
            'total_paragraphs': len(self.doc.paragraphs),
            'total_runs': 0,
            'font_issues': [],
            'bold_runs': [],
            'italic_runs': [],
            'all_compliant': True
        }

        for para_idx, paragraph in enumerate(self.doc.paragraphs, 1):
            for run_idx, run in enumerate(paragraph.runs, 1):
                results['total_runs'] += 1

                # Check font name
                font_name = run.font.name
                if font_name and font_name != 'Times New Roman':
                    issue = f"Para {para_idx}, Run {run_idx}: Font '{font_name}' (expected Times New Roman)"
                    results['font_issues'].append(issue)
                    results['all_compliant'] = False

                # Check font size
                font_size = run.font.size
                if font_size and font_size != Pt(12):
                    size_pt = font_size.pt if font_size else 'None'
                    issue = f"Para {para_idx}, Run {run_idx}: Size {size_pt}pt (expected 12pt)"
                    results['font_issues'].append(issue)
                    results['all_compliant'] = False

                # Check bold
                if run.font.bold:
                    issue = f"Para {para_idx}, Run {run_idx}: Bold text found"
                    results['bold_runs'].append(issue)
                    results['all_compliant'] = False

                # Check italic
                if run.font.italic:
                    issue = f"Para {para_idx}, Run {run_idx}: Italic text found"
                    results['italic_runs'].append(issue)
                    results['all_compliant'] = False

        return results

    def print_analysis(self, label: str, expected_pages: Optional[int] = None):
        """Print analysis results."""
        results = self.analyze_formatting()

        print(f"\n{label}")
        print(f"  - Pages: {results['page_count']}", end='')
        if expected_pages:
            if results['page_count'] == expected_pages:
                print(" ✓")
            else:
                print(f" ❌ (expected {expected_pages})")
                self.issues.append(f"Page count mismatch: {results['page_count']} (expected {expected_pages})")
        else:
            print()

        print(f"  - Paragraphs: {results['total_paragraphs']}")
        print(f"  - Runs: {results['total_runs']}")

        # Font issues
        if results['font_issues']:
            print(f"  - Font: ❌ ({len(results['font_issues'])} issues)")
            for issue in results['font_issues'][:5]:  # Show first 5
                print(f"      {issue}")
            if len(results['font_issues']) > 5:
                print(f"      ... and {len(results['font_issues']) - 5} more")
            self.issues.extend(results['font_issues'])
        else:
            print("  - Font: Times New Roman 12pt ✓")

        # Bold/Italic
        if results['bold_runs']:
            print(f"  - Bold: ❌ ({len(results['bold_runs'])} runs)")
            self.issues.extend(results['bold_runs'])
        else:
            print("  - Bold: None ✓")

        if results['italic_runs']:
            print(f"  - Italic: ❌ ({len(results['italic_runs'])} runs)")
            self.issues.extend(results['italic_runs'])
        else:
            print("  - Italic: None ✓")

        return results


def run_ocr_processing(pdf_path: str, output_path: str) -> bool:
    """
    Run OCR processing using LandingAI provider.

    Args:
        pdf_path: Path to input PDF
        output_path: Path for output DOCX

    Returns:
        True if successful
    """
    try:
        logger.info(f"Running OCR on {pdf_path}")

        # Import OCR provider
        from src.ocr import OCRProviderFactory
        from src.config import load_config

        # Load config
        config = load_config()

        # Get OCR provider (should be LandingAI based on config)
        provider = OCRProviderFactory.get_provider(config)

        logger.info(f"Using OCR provider: {provider.__class__.__name__}")

        # Process document
        provider.process_document(pdf_path, output_path)

        if os.path.exists(output_path):
            logger.info(f"OCR completed: {output_path}")
            return True
        else:
            logger.error(f"OCR failed: output file not created")
            return False

    except Exception as e:
        logger.error(f"OCR processing error: {e}", exc_info=True)
        return False


def run_translation(input_docx: str, output_docx: str, target_lang: str = 'en') -> bool:
    """
    Run GoogleTranslator on a DOCX file.

    Args:
        input_docx: Path to input DOCX
        output_docx: Path for output DOCX
        target_lang: Target language code

    Returns:
        True if successful
    """
    try:
        logger.info(f"Running translation on {input_docx}")

        # Build command to run GoogleTranslator with its own virtual environment
        translator_script = GOOGLE_TRANSLATOR_ROOT / "translate_document.py"
        translator_python = GOOGLE_TRANSLATOR_ROOT / "venv" / "bin" / "python3"

        # Check if GoogleTranslator has venv
        if not translator_python.exists():
            logger.error(f"GoogleTranslator venv not found at: {translator_python}")
            return False

        cmd = [
            str(translator_python),  # Use GoogleTranslator's Python
            str(translator_script),
            "--input", input_docx,
            "--output", output_docx,
            "--target", target_lang
        ]

        logger.debug(f"Running command: {' '.join(cmd)}")

        # Run translation
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            cwd=str(GOOGLE_TRANSLATOR_ROOT)
        )

        if result.returncode == 0:
            logger.info(f"Translation completed: {output_docx}")
            if result.stdout:
                logger.debug(f"Translation output: {result.stdout}")
            return True
        else:
            logger.error(f"Translation failed with code {result.returncode}")
            if result.stderr:
                logger.error(f"Translation stderr: {result.stderr}")
            if result.stdout:
                logger.error(f"Translation stdout: {result.stdout}")
            return False

    except Exception as e:
        logger.error(f"Translation error: {e}", exc_info=True)
        return False


def get_pdf_page_count(pdf_path: str) -> int:
    """Get page count from PDF file."""
    try:
        import PyPDF2
        with open(pdf_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            return len(pdf_reader.pages)
    except:
        # Fallback: try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                return len(pdf.pages)
        except:
            logger.warning("Could not determine PDF page count")
            return 0


def run_iteration(iteration: int, pdf_path: str, max_iterations: int = 10) -> Tuple[bool, List[str]]:
    """
    Run one iteration of the test.

    Args:
        iteration: Current iteration number (1-based)
        pdf_path: Path to PDF file to test
        max_iterations: Maximum number of iterations

    Returns:
        Tuple of (success, issues_found)
    """
    print(f"\n{'='*80}")
    print(f"ITERATION {iteration}/{max_iterations}")
    print(f"{'='*80}")

    # Generate timestamped filenames
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    pdf_name = Path(pdf_path).stem

    ocr_output = PROJECT_ROOT / f"{pdf_name}_ocr_test_{timestamp}.docx"
    translated_output = PROJECT_ROOT / f"{pdf_name}_translated_test_{timestamp}.docx"

    all_issues = []

    # Get original PDF page count
    original_pages = get_pdf_page_count(pdf_path)
    print(f"\nOriginal PDF: {Path(pdf_path).name} ({original_pages} pages)")

    # Step 1: OCR Processing
    print(f"\nStep 1: OCR Processing")
    print(f"  Input:  {Path(pdf_path).name}")
    print(f"  Output: {ocr_output.name}")

    if not run_ocr_processing(str(pdf_path), str(ocr_output)):
        print("  ❌ OCR FAILED")
        all_issues.append("OCR processing failed")
        return False, all_issues

    # Analyze OCR output
    ocr_analyzer = DocumentAnalyzer(str(ocr_output))
    ocr_results = ocr_analyzer.print_analysis(
        "  OCR Output Analysis:",
        expected_pages=original_pages
    )
    all_issues.extend(ocr_analyzer.issues)

    # Step 2: Translation
    print(f"\nStep 2: Translation")
    print(f"  Input:  {ocr_output.name}")
    print(f"  Output: {translated_output.name}")

    if not run_translation(str(ocr_output), str(translated_output)):
        print("  ❌ TRANSLATION FAILED")
        all_issues.append("Translation failed")
        return False, all_issues

    # Analyze translated output
    trans_analyzer = DocumentAnalyzer(str(translated_output))
    trans_results = trans_analyzer.print_analysis(
        "  Translation Output Analysis:",
        expected_pages=original_pages
    )
    all_issues.extend(trans_analyzer.issues)

    # Overall analysis
    print(f"\n{'='*80}")
    print("ITERATION SUMMARY")
    print(f"{'='*80}")

    if not all_issues:
        print("✓ ALL CRITERIA PASSED!")
        print(f"  - Page count: {trans_results['page_count']} (matches original)")
        print(f"  - Font: Times New Roman 12pt")
        print(f"  - No bold or italic")
        return True, []
    else:
        print(f"❌ FOUND {len(all_issues)} ISSUES:")

        # Categorize issues
        page_issues = [i for i in all_issues if 'page count' in i.lower()]
        font_issues = [i for i in all_issues if 'font' in i.lower() or 'size' in i.lower()]
        bold_issues = [i for i in all_issues if 'bold' in i.lower()]
        italic_issues = [i for i in all_issues if 'italic' in i.lower()]
        other_issues = [i for i in all_issues if i not in page_issues + font_issues + bold_issues + italic_issues]

        if page_issues:
            print(f"\n  Page Count Issues ({len(page_issues)}):")
            for issue in page_issues[:3]:
                print(f"    - {issue}")

        if font_issues:
            print(f"\n  Font Issues ({len(font_issues)}):")
            for issue in font_issues[:3]:
                print(f"    - {issue}")
            if len(font_issues) > 3:
                print(f"    ... and {len(font_issues) - 3} more")

        if bold_issues:
            print(f"\n  Bold Issues ({len(bold_issues)}):")
            for issue in bold_issues[:3]:
                print(f"    - {issue}")

        if italic_issues:
            print(f"\n  Italic Issues ({len(italic_issues)}):")
            for issue in italic_issues[:3]:
                print(f"    - {issue}")

        if other_issues:
            print(f"\n  Other Issues ({len(other_issues)}):")
            for issue in other_issues[:3]:
                print(f"    - {issue}")

        # Root cause analysis
        print(f"\n{'='*80}")
        print("ROOT CAUSE ANALYSIS")
        print(f"{'='*80}")

        if page_issues and ocr_results['page_count'] == original_pages and trans_results['page_count'] != original_pages:
            print("\n❌ PAGE COUNT ISSUE:")
            print(f"  Problem: Pages lost during translation")
            print(f"    OCR output: {ocr_results['page_count']} pages ✓")
            print(f"    Translated: {trans_results['page_count']} pages ❌")
            print(f"\n  Root Cause:")
            print(f"    GoogleTranslator is not preserving page breaks")
            print(f"\n  Location:")
            print(f"    GoogleTranslator/translator/core/translator.py")
            print(f"    Line ~728: Document creation without page break copying")
            print(f"\n  Fix Required:")
            print(f"    Copy paragraph_format.page_break_before from source to target")

        elif page_issues and ocr_results['page_count'] != original_pages:
            print("\n❌ PAGE COUNT ISSUE:")
            print(f"  Problem: Pages lost during OCR")
            print(f"    Original PDF: {original_pages} pages")
            print(f"    OCR output: {ocr_results['page_count']} pages ❌")
            print(f"\n  Root Cause:")
            print(f"    LandingAI provider not creating proper page breaks")
            print(f"\n  Location:")
            print(f"    EmailReader/src/convert_to_docx.py")
            print(f"    Function: convert_structured_to_docx()")
            print(f"\n  Fix Required:")
            print(f"    Verify page break creation in convert_structured_to_docx")

        if font_issues:
            print("\n❌ FONT ISSUES:")
            # Check which stage introduced the issue
            if not ocr_analyzer.issues:
                print(f"  Problem: Wrong fonts introduced during translation")
                print(f"    OCR output: All Times New Roman 12pt ✓")
                print(f"    Translated: Mixed fonts ❌")
                print(f"\n  Root Cause:")
                print(f"    GoogleTranslator not preserving font settings")
                print(f"\n  Location:")
                print(f"    GoogleTranslator/translator/core/translator.py")
                print(f"    Line ~728: Font copying logic")
            else:
                print(f"  Problem: Wrong fonts in OCR output")
                print(f"    OCR output: Mixed fonts ❌")
                print(f"\n  Root Cause:")
                print(f"    LandingAI provider not applying Times New Roman 12pt")
                print(f"\n  Location:")
                print(f"    EmailReader/src/convert_to_docx.py")
                print(f"    Function: convert_structured_to_docx()")
                print(f"    Lines 246-249, 268-272: Font application")

        if bold_issues or italic_issues:
            print("\n❌ BOLD/ITALIC ISSUES:")
            print(f"  Problem: Bold or italic formatting present")
            print(f"\n  Root Cause:")
            print(f"    Font bold/italic flags not being reset to False")
            print(f"\n  Location:")
            print(f"    EmailReader/src/convert_to_docx.py or GoogleTranslator")
            print(f"    Need to explicitly set run.font.bold = False, run.font.italic = False")

        return False, all_issues


def main():
    """Main test execution."""
    print("="*80)
    print("EMAILREADER FORMATTING VALIDATION TEST")
    print("Iterative Testing with Automatic Issue Detection")
    print("="*80)

    # Find test PDF
    test_pdf = None

    # Try Konnova.pdf first (if mentioned by user)
    konnova_pdf = INBOX_TEMP_DIR / "Konnova.pdf"
    if konnova_pdf.exists():
        test_pdf = konnova_pdf
        print(f"\nUsing test file: {test_pdf}")
    else:
        # Try PDF-scanned-rus-words.pdf from test_docs
        rus_pdf = TEST_DOCS_DIR / "PDF-scanned-rus-words.pdf"
        if rus_pdf.exists():
            test_pdf = rus_pdf
            print(f"\nUsing test file: {test_pdf}")
        else:
            # Try any PDF in test_docs
            pdf_files = list(TEST_DOCS_DIR.glob("*.pdf"))
            if pdf_files:
                test_pdf = pdf_files[0]
                print(f"\nUsing test file: {test_pdf}")
            else:
                print("\n❌ ERROR: No PDF files found for testing")
                print(f"  Looked in:")
                print(f"    - {INBOX_TEMP_DIR}/Konnova.pdf")
                print(f"    - {TEST_DOCS_DIR}/PDF-scanned-rus-words.pdf")
                print(f"    - {TEST_DOCS_DIR}/*.pdf")
                return 1

    # Run iterations
    max_iterations = 10
    for iteration in range(1, max_iterations + 1):
        success, issues = run_iteration(iteration, str(test_pdf), max_iterations)

        if success:
            print(f"\n{'='*80}")
            print(f"✓ SUCCESS! All criteria passed in iteration {iteration}/{max_iterations}")
            print(f"{'='*80}")
            return 0

        if iteration < max_iterations:
            print(f"\n{'='*80}")
            print(f"Continuing to iteration {iteration + 1}...")
            print(f"{'='*80}")
            # In a real scenario, we would apply fixes here
            # For this test, we're identifying issues that need manual fixes
        else:
            print(f"\n{'='*80}")
            print(f"❌ FAILED: Maximum iterations ({max_iterations}) reached")
            print(f"{'='*80}")
            print("\nIssues found require code fixes. Please review the root cause analysis above.")
            return 1


if __name__ == '__main__':
    sys.exit(main())
