"""
Test saving Azure paragraphs directly to DOCX
"""
from src.config import load_config
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from docx import Document
from docx.shared import Pt

# Load config
config = load_config()
azure_config = config['ocr']['azure']

# Create client
client = DocumentAnalysisClient(
    endpoint=azure_config['endpoint'],
    credential=AzureKeyCredential(azure_config['api_key'])
)

# Read PDF
pdf_path = "ФТР Артем Строкань EB1  (1) (2).pdf"
with open(pdf_path, 'rb') as f:
    pdf_bytes = f.read()

print(f"Processing {pdf_path}...")

# Analyze
poller = client.begin_analyze_document("prebuilt-read", document=pdf_bytes)
result = poller.result()

print(f"Paragraphs: {len(result.paragraphs)}")

# Create DOCX directly from Azure paragraphs
document = Document()

for page_num in range(1, len(result.pages) + 1):
    # Get paragraphs for this page
    page_paragraphs = [
        para
        for para in result.paragraphs
        if hasattr(para, 'bounding_regions') and
           para.bounding_regions and
           para.bounding_regions[0].page_number == page_num
    ]

    print(f"\nPage {page_num}: {len(page_paragraphs)} paragraphs")

    for i, para in enumerate(page_paragraphs):
        content = para.content
        has_newline = '\n' in content
        print(f"  Para {i+1}: {len(content)} chars, has_newline={has_newline}, preview={repr(content[:60])}")

        # Add to document
        p = document.add_paragraph(content)
        for run in p.runs:
            run.font.size = Pt(11)

output_path = "test_direct_paragraphs.docx"
document.save(output_path)
print(f"\nSaved to: {output_path}")
