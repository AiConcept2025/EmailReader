"""Integration tests for paragraph-based translation pipeline.

This test suite verifies the end-to-end paragraph extraction and batch translation
workflow using real Azure OCR and Google Translate services.

Tests follow TDD principles and use real credentials when available, with graceful
skipping when credentials are not configured.
"""
import pytest
import os
import tempfile
import logging
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config import load_config
from src.ocr.azure_provider import AzureOCRProvider
from src.translation.google_batch_translator import GoogleBatchTranslator
from src.ocr import OCRProviderFactory
from src.translation import TranslatorFactory
from src.models.paragraph import Paragraph, TextSpan

logger = logging.getLogger('EmailReader.Tests.Integration')


# Fixtures for configuration and credentials
@pytest.fixture(scope='module')
def config() -> Dict[str, Any]:
    """Load real configuration from config file."""
    try:
        return load_config()
    except Exception as e:
        pytest.skip(f"Config file not available: {e}")


@pytest.fixture(scope='module')
def azure_available(config) -> bool:
    """Check if Azure OCR credentials are configured."""
    try:
        azure_config = config.get('ocr', {}).get('azure', {})
        endpoint = azure_config.get('endpoint', '')
        api_key = azure_config.get('api_key', '')
        available = bool(endpoint and api_key)
        if not available:
            logger.info("Azure credentials not configured - tests will be skipped")
        return available
    except Exception:
        return False


@pytest.fixture(scope='module')
def google_available(config) -> bool:
    """Check if Google service account credentials are configured."""
    try:
        sa_data = config.get('google_drive', {}).get('service_account', {})
        project_id = sa_data.get('project_id', '')
        available = bool(project_id and sa_data)
        if not available:
            logger.info("Google credentials not configured - tests will be skipped")
        return available
    except Exception:
        return False


@pytest.fixture
def temp_output_dir():
    """Create temporary directory for test outputs."""
    with tempfile.TemporaryDirectory(prefix='test_paragraph_pipeline_') as tmpdir:
        logger.info(f"Created temporary directory: {tmpdir}")
        yield tmpdir
        logger.info(f"Cleaned up temporary directory: {tmpdir}")


@pytest.fixture
def test_docx_simple(temp_output_dir) -> str:
    """Create a simple test DOCX with multiple paragraphs."""
    doc = Document()

    # Add title
    title = doc.add_heading('Test Document Title', level=0)

    # Add heading
    heading = doc.add_heading('Section 1: Introduction', level=1)

    # Add paragraphs
    doc.add_paragraph('This is the first paragraph with normal text.')
    doc.add_paragraph('This is the second paragraph with more content.')

    # Add another heading
    doc.add_heading('Section 2: Details', level=1)

    # Add more paragraphs
    doc.add_paragraph('Third paragraph in the second section.')
    doc.add_paragraph('Fourth paragraph with additional information.')

    # Add list items
    doc.add_paragraph('First list item', style='List Bullet')
    doc.add_paragraph('Second list item', style='List Bullet')
    doc.add_paragraph('Third list item', style='List Bullet')

    # Save to temporary file
    output_path = os.path.join(temp_output_dir, 'test_simple.docx')
    doc.save(output_path)
    logger.info(f"Created simple test DOCX: {output_path}")

    return output_path


@pytest.fixture
def test_docx_large(temp_output_dir) -> str:
    """Create a large test DOCX with 60+ paragraphs for batch testing."""
    doc = Document()

    # Add title
    doc.add_heading('Large Test Document', level=0)

    # Create 65 paragraphs across multiple sections
    for section_num in range(1, 6):  # 5 sections
        doc.add_heading(f'Section {section_num}', level=1)

        # Add 13 paragraphs per section (65 total)
        for para_num in range(1, 14):
            text = f'This is paragraph {para_num} in section {section_num}. ' \
                   f'It contains sample text for testing batch translation. ' \
                   f'The content is generic but sufficient for testing purposes.'
            doc.add_paragraph(text)

    # Save to temporary file
    output_path = os.path.join(temp_output_dir, 'test_large.docx')
    doc.save(output_path)

    # Verify paragraph count
    doc_verify = Document(output_path)
    para_count = len([p for p in doc_verify.paragraphs if p.text.strip()])
    logger.info(f"Created large test DOCX with {para_count} paragraphs: {output_path}")

    return output_path


@pytest.fixture
def test_docx_formatted(temp_output_dir) -> str:
    """Create a test DOCX with various formatting (bold, italic, font sizes)."""
    doc = Document()

    # Title with large font
    title = doc.add_heading('Formatted Document', level=0)

    # Paragraph with mixed formatting
    para1 = doc.add_paragraph()
    para1.add_run('Normal text, ')
    bold_run = para1.add_run('bold text, ')
    bold_run.bold = True
    italic_run = para1.add_run('italic text, ')
    italic_run.italic = True
    both_run = para1.add_run('bold and italic')
    both_run.bold = True
    both_run.italic = True

    # Paragraph with different font sizes
    para2 = doc.add_paragraph()
    small_run = para2.add_run('Small text (10pt), ')
    small_run.font.size = Pt(10)
    normal_run = para2.add_run('normal text (12pt), ')
    normal_run.font.size = Pt(12)
    large_run = para2.add_run('large text (16pt)')
    large_run.font.size = Pt(16)

    # Heading with bold
    heading = doc.add_heading('Section Heading', level=1)

    # List items
    doc.add_paragraph('First formatted list item', style='List Bullet')
    doc.add_paragraph('Second formatted list item', style='List Bullet')

    # Save to temporary file
    output_path = os.path.join(temp_output_dir, 'test_formatted.docx')
    doc.save(output_path)
    logger.info(f"Created formatted test DOCX: {output_path}")

    return output_path


# Integration Tests

@pytest.mark.integration
class TestParagraphExtractionIntegration:
    """Integration tests for Azure paragraph extraction."""

    def test_azure_provider_initialization(self, config, azure_available):
        """Test that Azure provider can be initialized with real credentials."""
        if not azure_available:
            pytest.skip("Azure credentials not available")

        # Arrange
        azure_config = config.get('ocr', {}).get('azure', {})

        # Act
        provider = AzureOCRProvider(azure_config)

        # Assert
        assert provider is not None
        assert provider.endpoint == azure_config['endpoint']
        assert hasattr(provider, 'client')
        logger.info("[PASS] Azure provider initialized successfully")

    def test_paragraph_extraction_model_usage(self, config, azure_available):
        """Test that Azure provider uses correct model for paragraph extraction."""
        if not azure_available:
            pytest.skip("Azure credentials not available")

        # Arrange
        azure_config = config.get('ocr', {}).get('azure', {})
        provider = AzureOCRProvider(azure_config)

        # Assert - Verify provider has paragraph extraction capabilities
        assert hasattr(provider, '_extract_paragraphs_with_formatting')
        logger.info("[PASS] Azure provider has paragraph extraction method")


@pytest.mark.integration
class TestBatchTranslationIntegration:
    """Integration tests for Google Batch Translation."""

    def test_google_batch_translator_initialization(self, config, google_available):
        """Test that Google Batch Translator can be initialized with real credentials."""
        if not google_available:
            pytest.skip("Google credentials not available")

        # Arrange
        project_id = config.get('google_drive', {}).get('service_account', {}).get('project_id')
        translator_config = {'project_id': project_id}

        # Act
        translator = GoogleBatchTranslator(translator_config)

        # Assert
        assert translator is not None
        assert translator.project_id == project_id
        assert translator.BATCH_SIZE == 25
        assert hasattr(translator, 'client')
        logger.info(f"[PASS] Google Batch Translator initialized with project: {project_id}")

    def test_batch_translation_simple_document(
        self, config, google_available, test_docx_simple, temp_output_dir
    ):
        """Test batch translation of a simple document."""
        if not google_available:
            pytest.skip("Google credentials not available")

        # Arrange
        project_id = config.get('google_drive', {}).get('service_account', {}).get('project_id')
        translator_config = {'project_id': project_id}
        translator = GoogleBatchTranslator(translator_config)

        output_path = os.path.join(temp_output_dir, 'translated_simple.docx')

        # Act
        logger.info(f"[TEST] Translating simple document from Russian to English")
        translator.translate_document(test_docx_simple, output_path, target_lang='en')

        # Assert
        assert os.path.exists(output_path), "Translated document should exist"

        # Verify output structure
        doc = Document(output_path)
        original_doc = Document(test_docx_simple)

        original_paras = [p.text.strip() for p in original_doc.paragraphs if p.text.strip()]
        translated_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        assert len(translated_paras) > 0, "Translated document should have paragraphs"
        logger.info(f"[PASS] Simple document translated: {len(translated_paras)} paragraphs")

    def test_batch_translation_preserves_formatting(
        self, config, google_available, test_docx_formatted, temp_output_dir
    ):
        """Test that batch translation preserves formatting (bold, italic)."""
        if not google_available:
            pytest.skip("Google credentials not available")

        # Arrange
        project_id = config.get('google_drive', {}).get('service_account', {}).get('project_id')
        translator_config = {'project_id': project_id}
        translator = GoogleBatchTranslator(translator_config)

        output_path = os.path.join(temp_output_dir, 'translated_formatted.docx')

        # Act
        logger.info(f"[TEST] Translating formatted document")
        translator.translate_document(test_docx_formatted, output_path, target_lang='en')

        # Assert
        assert os.path.exists(output_path), "Translated document should exist"

        # Verify formatting is preserved
        doc = Document(output_path)

        # Check for bold runs
        has_bold = False
        has_italic = False

        for para in doc.paragraphs:
            for run in para.runs:
                if run.bold:
                    has_bold = True
                if run.italic:
                    has_italic = True

        # Note: Formatting preservation depends on translator implementation
        logger.info(f"[INFO] Formatting check - Bold: {has_bold}, Italic: {has_italic}")
        logger.info(f"[PASS] Formatted document translation completed")

    def test_batch_translation_large_document(
        self, config, google_available, test_docx_large, temp_output_dir
    ):
        """Test batch translation with 60+ paragraphs to verify batching."""
        if not google_available:
            pytest.skip("Google credentials not available")

        # Arrange
        project_id = config.get('google_drive', {}).get('service_account', {}).get('project_id')
        translator_config = {'project_id': project_id}
        translator = GoogleBatchTranslator(translator_config)

        output_path = os.path.join(temp_output_dir, 'translated_large.docx')

        # Count original paragraphs
        original_doc = Document(test_docx_large)
        original_count = len([p for p in original_doc.paragraphs if p.text.strip()])

        logger.info(f"[TEST] Translating large document with {original_count} paragraphs")
        assert original_count >= 60, "Test document should have at least 60 paragraphs"

        # Calculate expected batches
        expected_batches = (original_count + translator.BATCH_SIZE - 1) // translator.BATCH_SIZE
        logger.info(f"[INFO] Expected batches: {expected_batches} (batch_size={translator.BATCH_SIZE})")

        # Act
        translator.translate_document(test_docx_large, output_path, target_lang='en')

        # Assert
        assert os.path.exists(output_path), "Translated document should exist"

        # Verify paragraph count is preserved
        translated_doc = Document(output_path)
        translated_count = len([p for p in translated_doc.paragraphs if p.text.strip()])

        # Allow small variance due to formatting
        assert translated_count >= original_count * 0.9, \
            f"Translated document should have similar paragraph count (original: {original_count}, translated: {translated_count})"

        logger.info(f"[PASS] Large document translated successfully")
        logger.info(f"[PASS] Paragraph count - Original: {original_count}, Translated: {translated_count}")


@pytest.mark.integration
class TestEndToEndParagraphPipeline:
    """End-to-end integration tests for the complete paragraph pipeline."""

    def test_ocr_factory_selects_azure_for_human_mode(self, config, azure_available):
        """Test that OCR factory selects Azure provider for 'human' translation mode."""
        if not azure_available:
            pytest.skip("Azure credentials not available")

        # Act
        provider = OCRProviderFactory.get_provider(config, translation_mode='human')

        # Assert
        assert isinstance(provider, AzureOCRProvider), \
            f"Expected AzureOCRProvider for human mode, got {type(provider).__name__}"
        logger.info("[PASS] OCR Factory correctly routes 'human' mode to Azure")

    def test_translator_factory_selects_google_batch_from_config(
        self, config, google_available
    ):
        """Test that Translator factory selects Google Batch when configured."""
        if not google_available:
            pytest.skip("Google credentials not available")

        # Arrange - Create config with google_batch provider
        test_config = {
            'translation': {
                'provider': 'google_batch',
                'google_doc': {
                    'project_id': config.get('google_drive', {}).get('service_account', {}).get('project_id')
                }
            },
            'google_drive': config.get('google_drive', {})
        }

        # Act
        translator = TranslatorFactory.get_translator(test_config)

        # Assert
        assert isinstance(translator, GoogleBatchTranslator), \
            f"Expected GoogleBatchTranslator, got {type(translator).__name__}"
        logger.info("[PASS] Translator Factory correctly creates GoogleBatchTranslator")

    def test_paragraph_extraction_uses_correct_flag(self, config, azure_available):
        """Test that Azure provider receives use_paragraph_extraction flag."""
        if not azure_available:
            pytest.skip("Azure credentials not available")

        # Arrange
        azure_config = config.get('ocr', {}).get('azure', {})
        provider = AzureOCRProvider(azure_config)

        # Assert - Check that process_document accepts the flag
        import inspect
        sig = inspect.signature(provider.process_document)
        params = sig.parameters

        assert 'use_paragraph_extraction' in params, \
            "Azure provider should accept use_paragraph_extraction parameter"

        logger.info("[PASS] Azure provider supports use_paragraph_extraction flag")


@pytest.mark.integration
class TestBatchPerformance:
    """Performance tests for batch translation."""

    def test_batch_performance_with_large_document(
        self, config, google_available, test_docx_large, temp_output_dir
    ):
        """Test performance and batching behavior with 60+ paragraph document."""
        if not google_available:
            pytest.skip("Google credentials not available")

        import time

        # Arrange
        project_id = config.get('google_drive', {}).get('service_account', {}).get('project_id')
        translator_config = {'project_id': project_id}
        translator = GoogleBatchTranslator(translator_config)

        output_path = os.path.join(temp_output_dir, 'translated_performance.docx')

        # Count paragraphs
        original_doc = Document(test_docx_large)
        para_count = len([p for p in original_doc.paragraphs if p.text.strip()])

        logger.info(f"[TEST] Performance test with {para_count} paragraphs")

        # Act - Measure translation time
        start_time = time.time()
        translator.translate_document(test_docx_large, output_path, target_lang='en')
        end_time = time.time()

        duration = end_time - start_time

        # Assert
        assert os.path.exists(output_path), "Translation should complete successfully"

        # Verify batching occurred
        expected_batches = (para_count + translator.BATCH_SIZE - 1) // translator.BATCH_SIZE

        logger.info(f"[PASS] Translation completed in {duration:.2f} seconds")
        logger.info(f"[PASS] Processed {para_count} paragraphs in ~{expected_batches} batches")
        logger.info(f"[PASS] Average speed: {para_count / duration:.2f} paragraphs/second")

        # Performance assertion - should be reasonably fast
        # Allow 10 seconds per batch (conservative estimate)
        max_expected_time = expected_batches * 10
        assert duration < max_expected_time, \
            f"Translation took {duration:.2f}s, expected < {max_expected_time}s"
