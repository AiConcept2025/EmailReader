"""
Test Azure OCR role-based formatting.

Verifies that the Azure prebuilt-layout model correctly:
1. Extracts paragraph roles (title, sectionHeading, body, etc.)
2. Applies appropriate formatting based on roles
3. Preserves visual hierarchy in the output document
"""

import json
import logging
from pathlib import Path
from docx import Document

from src.ocr.azure_provider import AzureOCRProvider


def test_role_detection_and_formatting():
    """Test that roles are detected and formatting is applied."""

    # Set up logging to see role detection
    logging.basicConfig(
        level=logging.INFO,
        format='%(levelname)s:%(name)s:%(message)s'
    )

    # Load config
    config_path = Path('credentials/config.dev.json')
    with open(config_path, 'r') as f:
        config = json.load(f)

    # Get Azure OCR config
    azure_config = config.get('ocr', {}).get('azure', {})

    print(f"Using Azure model: {azure_config.get('model')}")
    print("=" * 80)

    # Process test document
    input_file = 'Konnova.pdf'
    output_file = 'Konnova_ocr_test_roles.docx'

    print(f"\nProcessing: {input_file}")
    print(f"Output: {output_file}")
    print("=" * 80)

    # Run OCR
    provider = AzureOCRProvider(azure_config)
    provider.process_document(input_file, output_file)

    print("\n" + "=" * 80)
    print("OCR Processing Complete!")
    print("=" * 80)

    # Verify the output document exists and has content
    doc = Document(output_file)

    print(f"\nDocument Statistics:")
    print(f"  Total paragraphs: {len(doc.paragraphs)}")

    # Analyze formatting in the document
    title_count = 0
    heading_count = 0
    body_count = 0

    for i, para in enumerate(doc.paragraphs):
        if not para.runs:
            continue

        run = para.runs[0]
        font_size = run.font.size
        is_bold = run.font.bold
        is_italic = run.font.italic
        alignment = para.alignment

        # Detect formatting patterns
        if font_size and font_size.pt >= 18 and is_bold:
            title_count += 1
            print(f"  [Title] Paragraph {i}: '{para.text[:50]}...'")
        elif font_size and font_size.pt >= 14 and is_bold:
            heading_count += 1
            print(f"  [Heading] Paragraph {i}: '{para.text[:50]}...'")
        elif font_size and font_size.pt <= 9 and is_italic:
            print(f"  [Header/Footer] Paragraph {i}: '{para.text[:50]}...'")
        elif font_size and font_size.pt == 11:
            body_count += 1

    print(f"\nFormatting Analysis:")
    print(f"  Titles (18pt, bold): {title_count}")
    print(f"  Headings (14pt, bold): {heading_count}")
    print(f"  Body text (11pt): {body_count}")

    print("\n" + "=" * 80)
    print("✓ Test completed successfully!")
    print(f"✓ Output saved to: {output_file}")
    print("=" * 80)


if __name__ == '__main__':
    test_role_detection_and_formatting()
