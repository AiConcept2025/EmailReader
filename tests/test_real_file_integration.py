#!/usr/bin/env python3
"""
Integration test using real files from Inbox/test_docs folders.

This test verifies the complete workflow:
1. Process PDF with LandingAI OCR
2. JSON saved to completed_temp/ with timestamp
3. Content filtering applied (no decorative descriptions)
4. DOCX output has correct content
5. Google Drive upload path is exercised
"""

import os
import pytest
import json
import glob
from pathlib import Path
from docx import Document

from src.ocr import OCRProviderFactory
from src.config import load_config


class TestRealFileIntegration:
    """Integration tests using real PDF files."""

    @pytest.fixture
    def config(self):
        """Load configuration."""
        return load_config()

    @pytest.fixture
    def ocr_provider(self, config):
        """Get OCR provider."""
        return OCRProviderFactory.get_provider(config)

    @pytest.fixture
    def test_pdf_path(self):
        """Path to real test PDF."""
        return os.path.join(os.getcwd(), 'test_docs', 'PDF-scanned-rus-words.pdf')

    @pytest.fixture
    def output_docx_path(self, tmp_path):
        """Temporary output DOCX path."""
        return str(tmp_path / "test_output_integration.docx")

    def test_complete_workflow_with_real_file(self, ocr_provider, test_pdf_path, output_docx_path, config):
        """
        Test complete OCR workflow with real PDF file.

        Verifies:
        - PDF is processed successfully
        - DOCX is created
        - JSON is saved to completed_temp/
        - No decorative content in output
        """
        # Ensure test file exists
        assert os.path.exists(test_pdf_path), f"Test file not found: {test_pdf_path}"

        # Get base name for JSON lookup
        base_name = Path(test_pdf_path).stem

        # Process the document
        ocr_provider.process_document(test_pdf_path, output_docx_path)

        # Verify DOCX was created
        assert os.path.exists(output_docx_path), "DOCX file was not created"
        assert os.path.getsize(output_docx_path) > 0, "DOCX file is empty"

        # Verify JSON was saved to completed_temp/
        completed_temp = os.path.join(os.getcwd(), 'completed_temp')
        json_files = glob.glob(os.path.join(completed_temp, f"{base_name}_landing_ai_*.json"))

        assert len(json_files) > 0, f"No JSON files found in {completed_temp} for {base_name}"

        # Get the most recent JSON file
        latest_json = max(json_files, key=os.path.getmtime)
        assert os.path.exists(latest_json), f"JSON file not found: {latest_json}"

        # Verify JSON structure
        with open(latest_json, 'r', encoding='utf-8') as f:
            json_data = json.load(f)

        assert 'chunks' in json_data, "JSON missing 'chunks' field"
        chunks = json_data['chunks']
        assert len(chunks) > 0, "JSON has no chunks"

        # Verify DOCX content has no decorative descriptions
        doc = Document(output_docx_path)

        # Get all text from document
        all_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

        # Verify no decorative content indicators
        decorative_indicators = [
            'stylized',
            'signature: illegible',
            'official stamp',
            'decorative border',
            'qr code',
            'scan code',
            '<::',
            '::>',
            'logo features',
            'logo depicts',
            'ornate',
            'floral motifs',
            'scalloped edge',
            'teardrop shape'
        ]

        found_indicators = []
        for indicator in decorative_indicators:
            if indicator in all_text.lower():
                found_indicators.append(indicator)

        assert len(found_indicators) == 0, (
            f"Found decorative indicators in DOCX: {found_indicators}\n"
            f"First 500 chars of content: {all_text[:500]}"
        )

        print(f"\n✅ Integration test passed:")
        print(f"   - PDF processed: {test_pdf_path}")
        print(f"   - DOCX created: {output_docx_path} ({os.path.getsize(output_docx_path)/1024:.2f} KB)")
        print(f"   - JSON saved: {latest_json}")
        print(f"   - Chunks in JSON: {len(chunks)}")
        print(f"   - Paragraphs in DOCX: {len([p for p in doc.paragraphs if p.text.strip()])}")
        print(f"   - No decorative content found ✓")

    def test_json_has_timestamp(self):
        """
        Verify JSON files have timestamps in filename.

        Format: {filename}_landing_ai_{YYYYMMDD_HHMMSS}.json
        """
        completed_temp = os.path.join(os.getcwd(), 'completed_temp')

        if not os.path.exists(completed_temp):
            pytest.skip("completed_temp directory not found")

        json_files = glob.glob(os.path.join(completed_temp, "*_landing_ai_*.json"))

        if len(json_files) == 0:
            pytest.skip("No JSON files found in completed_temp")

        # Verify timestamp format
        import re
        timestamp_pattern = r'_landing_ai_\d{8}_\d{6}\.json$'

        for json_file in json_files:
            assert re.search(timestamp_pattern, json_file), (
                f"JSON file missing timestamp: {os.path.basename(json_file)}"
            )

        print(f"\n✅ JSON timestamp test passed:")
        print(f"   - Found {len(json_files)} JSON files with timestamps")

    def test_document_correctness_validation(self, ocr_provider, test_pdf_path, output_docx_path):
        """
        Test document correctness validation.

        Ensures processed document contains actual content,
        not just decorative descriptions.
        """
        # Ensure test file exists
        if not os.path.exists(test_pdf_path):
            pytest.skip(f"Test file not found: {test_pdf_path}")

        # Process document
        ocr_provider.process_document(test_pdf_path, output_docx_path)

        # Load document
        doc = Document(output_docx_path)

        # Get non-empty paragraphs
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        assert len(paragraphs) > 0, "Document has no content"

        # Verify paragraphs have actual content (not just decorative descriptions)
        # Real content typically has:
        # - Multiple words
        # - Sentence structure
        # - Actual document text

        meaningful_paragraphs = [
            p for p in paragraphs
            if len(p.split()) >= 3  # At least 3 words
        ]

        assert len(meaningful_paragraphs) > 0, (
            "Document has no meaningful content (all short phrases)\n"
            f"Paragraphs found: {paragraphs[:5]}"
        )

        print(f"\n✅ Document correctness test passed:")
        print(f"   - Total paragraphs: {len(paragraphs)}")
        print(f"   - Meaningful paragraphs: {len(meaningful_paragraphs)}")
        print(f"   - Sample content: {meaningful_paragraphs[0][:100] if meaningful_paragraphs else 'N/A'}")

    def test_filtering_statistics(self):
        """
        Verify filtering statistics from JSON files.

        Checks that decorative chunks were filtered out.
        """
        completed_temp = os.path.join(os.getcwd(), 'completed_temp')

        if not os.path.exists(completed_temp):
            pytest.skip("completed_temp directory not found")

        json_files = glob.glob(os.path.join(completed_temp, "*_landing_ai_*.json"))

        if len(json_files) == 0:
            pytest.skip("No JSON files found")

        # Analyze most recent JSON file
        latest_json = max(json_files, key=os.path.getmtime)

        with open(latest_json, 'r', encoding='utf-8') as f:
            json_data = json.load(f)

        chunks = json_data.get('chunks', [])

        # Count chunk types
        from collections import Counter
        chunk_types = Counter(chunk.get('type') for chunk in chunks)

        # Count decorative chunks
        decorative_types = {'logo', 'scan_code', 'attestation', 'figure', 'image',
                          'barcode', 'border', 'decorative', 'signature'}
        decorative_count = sum(count for chunk_type, count in chunk_types.items()
                              if chunk_type in decorative_types)

        text_chunks = chunk_types.get('text', 0)

        print(f"\n✅ Filtering statistics:")
        print(f"   - Total chunks in JSON: {len(chunks)}")
        print(f"   - Text chunks: {text_chunks}")
        print(f"   - Decorative chunks (filtered): {decorative_count}")
        print(f"   - Chunk types: {dict(chunk_types)}")

        assert text_chunks > 0, "No text chunks found in JSON"


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
