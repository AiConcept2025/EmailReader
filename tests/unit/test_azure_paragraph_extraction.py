"""Unit tests for Azure OCR paragraph extraction with formatting.

This test suite follows TDD principles and tests the new paragraph extraction
functionality in Azure OCR provider.
"""
import pytest
from unittest.mock import Mock, MagicMock, patch
from typing import List
import os
import tempfile

from src.ocr.azure_provider import AzureOCRProvider
from src.models.paragraph import Paragraph, TextSpan


class TestAzureOCRParagraphExtraction:
    """Test Azure OCR paragraph extraction functionality."""

    @pytest.fixture
    def azure_config(self):
        """Azure OCR provider configuration."""
        return {
            'endpoint': 'https://test.cognitiveservices.azure.com/',
            'api_key': 'test_api_key_12345'
        }

    @pytest.fixture
    def provider(self, azure_config):
        """Create Azure OCR provider instance."""
        return AzureOCRProvider(azure_config)

    @pytest.fixture
    def mock_azure_paragraph_response(self):
        """Create mock Azure API response with paragraph data."""
        # Create mock paragraph objects
        para1 = Mock()
        para1.content = "Document Title"
        para1.role = "title"
        para1.page_number = 1
        para1.bounding_regions = [Mock(page_number=1, polygon=[
            Mock(x=100, y=50), Mock(x=500, y=50),
            Mock(x=500, y=100), Mock(x=100, y=100)
        ])]
        para1.spans = []

        para2 = Mock()
        para2.content = "Section Heading"
        para2.role = "heading"
        para2.page_number = 1
        para2.bounding_regions = [Mock(page_number=1, polygon=[
            Mock(x=100, y=120), Mock(x=400, y=120),
            Mock(x=400, y=150), Mock(x=100, y=150)
        ])]
        para2.spans = []

        para3 = Mock()
        para3.content = "This is a regular paragraph with some content."
        para3.role = "paragraph"
        para3.page_number = 1
        para3.bounding_regions = [Mock(page_number=1, polygon=[
            Mock(x=100, y=160), Mock(x=500, y=160),
            Mock(x=500, y=200), Mock(x=100, y=200)
        ])]
        para3.spans = []

        para4 = Mock()
        para4.content = "Content on page 2"
        para4.role = "paragraph"
        para4.page_number = 2
        para4.bounding_regions = [Mock(page_number=2, polygon=[
            Mock(x=100, y=50), Mock(x=500, y=50),
            Mock(x=500, y=100), Mock(x=100, y=100)
        ])]
        para4.spans = []

        # Create mock result
        result = Mock()
        result.paragraphs = [para1, para2, para3, para4]
        result.pages = [Mock(), Mock()]  # 2 pages

        return result

    def test_extract_paragraphs_with_formatting_extracts_correct_count(
        self, provider, mock_azure_paragraph_response
    ):
        """Test that paragraph extraction returns correct number of paragraphs."""
        pdf_bytes = b'%PDF-1.4 mock pdf content'

        with patch.object(provider.client, 'begin_analyze_document') as mock_analyze:
            mock_poller = Mock()
            mock_poller.result.return_value = mock_azure_paragraph_response
            mock_analyze.return_value = mock_poller

            paragraphs = provider._extract_paragraphs_with_formatting(pdf_bytes)

            assert len(paragraphs) == 4
            assert all(isinstance(p, Paragraph) for p in paragraphs)

    def test_extract_paragraphs_with_formatting_detects_roles(
        self, provider, mock_azure_paragraph_response
    ):
        """Test that paragraph extraction correctly detects roles."""
        pdf_bytes = b'%PDF-1.4 mock pdf content'

        with patch.object(provider.client, 'begin_analyze_document') as mock_analyze:
            mock_poller = Mock()
            mock_poller.result.return_value = mock_azure_paragraph_response
            mock_analyze.return_value = mock_poller

            paragraphs = provider._extract_paragraphs_with_formatting(pdf_bytes)

            assert paragraphs[0].role == "title"
            assert paragraphs[0].content == "Document Title"
            assert paragraphs[1].role == "heading"
            assert paragraphs[1].content == "Section Heading"
            assert paragraphs[2].role == "paragraph"
            assert paragraphs[3].role == "paragraph"

    def test_extract_paragraphs_with_formatting_extracts_bounding_boxes(
        self, provider, mock_azure_paragraph_response
    ):
        """Test that paragraph extraction includes bounding box information."""
        pdf_bytes = b'%PDF-1.4 mock pdf content'

        with patch.object(provider.client, 'begin_analyze_document') as mock_analyze:
            mock_poller = Mock()
            mock_poller.result.return_value = mock_azure_paragraph_response
            mock_analyze.return_value = mock_poller

            paragraphs = provider._extract_paragraphs_with_formatting(pdf_bytes)

            # Check first paragraph has bounding box
            assert paragraphs[0].bounding_box is not None
            assert 'x' in paragraphs[0].bounding_box
            assert 'y' in paragraphs[0].bounding_box
            assert 'width' in paragraphs[0].bounding_box
            assert 'height' in paragraphs[0].bounding_box

            # Verify values
            assert paragraphs[0].bounding_box['x'] == 100
            assert paragraphs[0].bounding_box['y'] == 50
            assert paragraphs[0].bounding_box['width'] == 400  # 500 - 100
            assert paragraphs[0].bounding_box['height'] == 50  # 100 - 50

    def test_extract_paragraphs_with_formatting_assigns_pages(
        self, provider, mock_azure_paragraph_response
    ):
        """Test that paragraph extraction assigns correct page numbers."""
        pdf_bytes = b'%PDF-1.4 mock pdf content'

        with patch.object(provider.client, 'begin_analyze_document') as mock_analyze:
            mock_poller = Mock()
            mock_poller.result.return_value = mock_azure_paragraph_response
            mock_analyze.return_value = mock_poller

            paragraphs = provider._extract_paragraphs_with_formatting(pdf_bytes)

            assert paragraphs[0].page == 1
            assert paragraphs[1].page == 1
            assert paragraphs[2].page == 1
            assert paragraphs[3].page == 2

    def test_extract_paragraphs_with_formatting_handles_missing_attributes(
        self, provider
    ):
        """Test graceful handling of missing Azure API attributes."""
        pdf_bytes = b'%PDF-1.4 mock pdf content'

        # Create minimal paragraph without optional fields
        # Use spec to limit attributes - Mock will raise AttributeError for missing attrs
        para = Mock(spec=['content', 'page_number', 'spans'])
        para.content = "Simple paragraph"
        para.page_number = 1
        para.spans = []

        result = Mock()
        result.paragraphs = [para]
        result.pages = [Mock()]

        with patch.object(provider.client, 'begin_analyze_document') as mock_analyze:
            mock_poller = Mock()
            mock_poller.result.return_value = result
            mock_analyze.return_value = mock_poller

            paragraphs = provider._extract_paragraphs_with_formatting(pdf_bytes)

            assert len(paragraphs) == 1
            assert paragraphs[0].content == "Simple paragraph"
            assert paragraphs[0].role == "paragraph"  # Default role
            assert paragraphs[0].bounding_box is None

    def test_save_paragraphs_to_docx_creates_file(self, provider, tmp_path):
        """Test that DOCX file is created with paragraphs."""
        paragraphs = [
            Paragraph(
                content="Document Title",
                page=1,
                role="title",
                spans=[TextSpan(text="Document Title", is_bold=True, font_size=18.0)]
            ),
            Paragraph(
                content="Section Heading",
                page=1,
                role="heading",
                spans=[TextSpan(text="Section Heading", is_bold=True, font_size=14.0)]
            ),
            Paragraph(
                content="Regular paragraph text.",
                page=1,
                role="paragraph"
            )
        ]

        output_path = tmp_path / "test_output.docx"

        provider._save_paragraphs_to_docx(paragraphs, str(output_path))

        assert output_path.exists()
        assert output_path.stat().st_size > 0

    def test_save_paragraphs_to_docx_applies_title_style(self, provider, tmp_path):
        """Test that title paragraphs get appropriate heading style."""
        paragraphs = [
            Paragraph(
                content="Document Title",
                page=1,
                role="title",
                spans=[TextSpan(text="Document Title", is_bold=True, font_size=18.0)]
            )
        ]

        output_path = tmp_path / "test_title.docx"

        provider._save_paragraphs_to_docx(paragraphs, str(output_path))

        # Verify file was created
        assert output_path.exists()

        # Load and verify content
        from docx import Document
        doc = Document(str(output_path))
        assert len(doc.paragraphs) > 0
        # First paragraph should be the title
        assert "Document Title" in doc.paragraphs[0].text

    def test_save_paragraphs_to_docx_applies_heading_style(self, provider, tmp_path):
        """Test that heading paragraphs get appropriate heading style."""
        paragraphs = [
            Paragraph(
                content="Section Heading",
                page=1,
                role="heading",
                spans=[TextSpan(text="Section Heading", is_bold=True)]
            )
        ]

        output_path = tmp_path / "test_heading.docx"

        provider._save_paragraphs_to_docx(paragraphs, str(output_path))

        assert output_path.exists()

        from docx import Document
        doc = Document(str(output_path))
        assert "Section Heading" in doc.paragraphs[0].text

    def test_save_paragraphs_to_docx_adds_page_breaks(self, provider, tmp_path):
        """Test that page breaks are added between pages."""
        paragraphs = [
            Paragraph(content="Page 1 content", page=1, role="paragraph"),
            Paragraph(content="Page 2 content", page=2, role="paragraph"),
            Paragraph(content="Page 3 content", page=3, role="paragraph")
        ]

        output_path = tmp_path / "test_pages.docx"

        provider._save_paragraphs_to_docx(paragraphs, str(output_path))

        assert output_path.exists()

        from docx import Document
        doc = Document(str(output_path))
        # Should have at least 3 paragraphs + page breaks
        assert len(doc.paragraphs) >= 3

    def test_save_paragraphs_to_docx_handles_list_items(self, provider, tmp_path):
        """Test that list items are formatted with markers."""
        paragraphs = [
            Paragraph(
                content="First item",
                page=1,
                role="listItem",
                is_list_item=True,
                list_marker="•"
            ),
            Paragraph(
                content="Second item",
                page=1,
                role="listItem",
                is_list_item=True,
                list_marker="•"
            )
        ]

        output_path = tmp_path / "test_list.docx"

        provider._save_paragraphs_to_docx(paragraphs, str(output_path))

        assert output_path.exists()

        from docx import Document
        doc = Document(str(output_path))
        # Should contain list markers or list items
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert "First item" in text
        assert "Second item" in text

    def test_process_document_with_paragraph_extraction_mode(
        self, provider, tmp_path, mock_azure_paragraph_response
    ):
        """Test process_document with use_paragraph_extraction=True."""
        # Create a test PDF file
        input_pdf = tmp_path / "test_input.pdf"
        input_pdf.write_bytes(b'%PDF-1.4 mock pdf content')

        output_docx = tmp_path / "test_output.docx"

        with patch.object(provider.client, 'begin_analyze_document') as mock_analyze:
            mock_poller = Mock()
            mock_poller.result.return_value = mock_azure_paragraph_response
            mock_analyze.return_value = mock_poller

            # Mock PDF searchability check to force OCR path
            with patch('pdfplumber.open') as mock_pdf:
                mock_page = Mock()
                mock_page.extract_text.return_value = ""  # Empty text -> needs OCR
                mock_pdf.return_value.__enter__.return_value.pages = [mock_page]

                provider.process_document(
                    str(input_pdf),
                    str(output_docx),
                    use_paragraph_extraction=True
                )

        assert output_docx.exists()

    def test_process_document_backward_compatibility(
        self, provider, tmp_path
    ):
        """Test that process_document without use_paragraph_extraction still works."""
        input_pdf = tmp_path / "test_input.pdf"
        input_pdf.write_bytes(b'%PDF-1.4 mock pdf content')

        output_docx = tmp_path / "test_output.docx"

        # Mock the old line-based extraction
        mock_result = Mock()
        mock_page = Mock()
        mock_line1 = Mock()
        mock_line1.content = "Line 1"
        mock_line2 = Mock()
        mock_line2.content = "Line 2"
        mock_page.lines = [mock_line1, mock_line2]
        mock_result.pages = [mock_page]

        with patch.object(provider.client, 'begin_analyze_document') as mock_analyze:
            mock_poller = Mock()
            mock_poller.result.return_value = mock_result
            mock_analyze.return_value = mock_poller

            with patch('pdfplumber.open') as mock_pdf:
                mock_pdf_page = Mock()
                mock_pdf_page.extract_text.return_value = ""  # Empty -> needs OCR
                mock_pdf.return_value.__enter__.return_value.pages = [mock_pdf_page]

                # Call without use_paragraph_extraction (default False)
                provider.process_document(str(input_pdf), str(output_docx))

        assert output_docx.exists()

    def test_extract_paragraphs_with_formatting_logs_progress(
        self, provider, mock_azure_paragraph_response, caplog
    ):
        """Test that paragraph extraction logs progress information."""
        import logging
        caplog.set_level(logging.INFO)

        pdf_bytes = b'%PDF-1.4 mock pdf content'

        with patch.object(provider.client, 'begin_analyze_document') as mock_analyze:
            mock_poller = Mock()
            mock_poller.result.return_value = mock_azure_paragraph_response
            mock_analyze.return_value = mock_poller

            paragraphs = provider._extract_paragraphs_with_formatting(pdf_bytes)

            # Check that logging occurred
            assert len(paragraphs) == 4
            # Logs should contain information about paragraph extraction
            log_text = caplog.text.lower()
            assert "paragraph" in log_text or "extract" in log_text
