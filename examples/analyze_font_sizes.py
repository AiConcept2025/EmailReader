"""
Analyze font sizes in a DOCX document
"""
from docx import Document
from collections import Counter
import sys

if len(sys.argv) < 2:
    print("Usage: python analyze_font_sizes.py <docx_file>")
    sys.exit(1)

doc_path = sys.argv[1]
doc = Document(doc_path)

font_sizes = []
text_samples = {}  # Store sample text for each font size

# Extract from paragraphs
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.font.size:
            size_pt = run.font.size.pt
            font_sizes.append(size_pt)
            if size_pt not in text_samples and run.text.strip():
                text_samples[size_pt] = run.text.strip()[:50]

# Extract from tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.size:
                        size_pt = run.font.size.pt
                        font_sizes.append(size_pt)
                        if size_pt not in text_samples and run.text.strip():
                            text_samples[size_pt] = run.text.strip()[:50]

# Count occurrences
size_counts = Counter(font_sizes)

print(f"\nDocument: {doc_path}")
print(f"Total font size samples: {len(font_sizes)}")
print("\nFont size distribution:")
print("-" * 80)
print(f"{'Size (pt)':<12} {'Count':<10} {'%':<8} {'Sample Text'}")
print("-" * 80)

for size, count in sorted(size_counts.items()):
    pct = (count / len(font_sizes)) * 100
    sample = text_samples.get(size, "")
    print(f"{size:<12.1f} {count:<10} {pct:<7.1f}% {sample}")

print("-" * 80)
print(f"\nUnique font sizes: {sorted(size_counts.keys())}")
