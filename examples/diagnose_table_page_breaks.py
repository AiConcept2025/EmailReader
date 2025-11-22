#!/usr/bin/env python3
"""
Diagnostic script to compare page breaks in table cells between OCR and translated outputs.
"""

import glob
from pathlib import Path
from docx import Document

def analyze_table_page_breaks(docx_path: str, label: str):
    """Analyze page breaks in tables."""
    print(f"\n{'='*80}")
    print(f"{label}: {Path(docx_path).name}")
    print('='*80)

    doc = Document(docx_path)

    # Count paragraphs before tables
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)

    print(f"Document has {para_count} paragraphs and {table_count} tables")

    # Analyze regular paragraphs with page breaks
    page_breaks_found = []
    for idx, para in enumerate(doc.paragraphs):
        if para.paragraph_format.page_break_before:
            page_breaks_found.append(f"Para {idx}")

    print(f"\nPage breaks in paragraphs: {len(page_breaks_found)}")
    if page_breaks_found:
        print(f"  Locations: {', '.join(page_breaks_found)}")

    # Analyze tables
    if doc.tables:
        print(f"\n{'Table Analysis':-^80}")
        for table_idx, table in enumerate(doc.tables):
            print(f"\nTable {table_idx}: {len(table.rows)} rows x {len(table.columns)} cols")

            # Check each cell
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_paras = cell.paragraphs
                    print(f"  Row {row_idx}, Col {col_idx}: {len(cell_paras)} paragraphs")

                    for para_idx, para in enumerate(cell_paras):
                        text_preview = para.text[:50] if para.text else "[empty]"
                        has_page_break = para.paragraph_format.page_break_before

                        print(f"    Para {para_idx}: page_break={has_page_break} | Text: {text_preview}")

                        if has_page_break:
                            print(f"      ⚠️  PAGE BREAK FOUND HERE!")

    # Count total pages (page breaks + 1)
    total_page_breaks = len(page_breaks_found)

    # Check for page breaks in tables
    table_page_breaks = 0
    if doc.tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.paragraph_format.page_break_before:
                            table_page_breaks += 1

    total_page_breaks += table_page_breaks
    total_pages = total_page_breaks + 1

    print(f"\n{'Summary':-^80}")
    print(f"Regular paragraph page breaks: {len(page_breaks_found)}")
    print(f"Table cell page breaks: {table_page_breaks}")
    print(f"Total page breaks: {total_page_breaks}")
    print(f"Total pages: {total_pages}")

    return total_pages


if __name__ == "__main__":
    # Find most recent OCR and translated test outputs
    project_root = Path(__file__).parent

    ocr_files = sorted(glob.glob(str(project_root / "PDF-scanned-rus-words_ocr_test_*.docx")))
    translated_files = sorted(glob.glob(str(project_root / "PDF-scanned-rus-words_translated_test_*.docx")))

    if not ocr_files or not translated_files:
        print("ERROR: No test output files found!")
        print("Please run test_formatting_iterative.py first")
        exit(1)

    # Use most recent files
    latest_ocr = ocr_files[-1]
    latest_translated = translated_files[-1]

    # Analyze both files
    ocr_pages = analyze_table_page_breaks(latest_ocr, "OCR OUTPUT")
    translated_pages = analyze_table_page_breaks(latest_translated, "TRANSLATED OUTPUT")

    print("\n" + "="*80)
    print("COMPARISON")
    print("="*80)
    print(f"OCR output: {ocr_pages} pages")
    print(f"Translated output: {translated_pages} pages")

    if ocr_pages == translated_pages:
        print("✅ Page count matches!")
    else:
        print(f"❌ Page count mismatch: {ocr_pages} → {translated_pages} (lost {ocr_pages - translated_pages} page)")
