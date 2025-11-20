#!/usr/bin/env python3
"""
Test script to verify paragraph filtering implementation.
Tests the new dual-output OCR pipeline with the Russian Tennis letter.
"""

import os
import logging
from src.config import load_config
from src.ocr import OCRProviderFactory

# Configure logging to both file and console
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/test_paragraph_filtering.log'),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger('TestParagraphFiltering')

def main():
    """Test the paragraph filtering with a real PDF."""

    # Input and output paths
    input_pdf = "ФТР Артем Строкань EB1  (1) (2).pdf"
    output_clean = "test_output_clean.docx"
    output_verification = "test_output_clean_ocr_verification.docx"

    logger.info("=" * 80)
    logger.info("PARAGRAPH FILTERING TEST")
    logger.info("=" * 80)
    logger.info("Input PDF: %s", input_pdf)
    logger.info("Expected outputs:")
    logger.info("  - Clean document: %s", output_clean)
    logger.info("  - Verification document: %s", output_verification)
    logger.info("")

    # Check input exists
    if not os.path.exists(input_pdf):
        logger.error("Input PDF not found: %s", input_pdf)
        return 1

    file_size_mb = os.path.getsize(input_pdf) / (1024 * 1024)
    logger.info("Input file size: %.2f MB", file_size_mb)

    # Load configuration
    logger.info("Loading configuration...")
    config = load_config()

    # Get OCR provider
    logger.info("Initializing OCR provider...")
    ocr_provider = OCRProviderFactory.get_provider(config)
    logger.info("OCR Provider: %s", type(ocr_provider).__name__)

    # Process document
    logger.info("")
    logger.info("Processing document with paragraph filtering...")
    logger.info("-" * 80)

    try:
        ocr_provider.process_document(input_pdf, output_clean)
        logger.info("-" * 80)
        logger.info("")
        logger.info("Processing completed successfully!")

        # Verify outputs exist
        logger.info("")
        logger.info("Verifying outputs...")

        if os.path.exists(output_clean):
            clean_size_kb = os.path.getsize(output_clean) / 1024
            logger.info("✓ Clean document created: %s (%.2f KB)", output_clean, clean_size_kb)
        else:
            logger.error("✗ Clean document NOT created")
            return 1

        if os.path.exists(output_verification):
            verification_size_kb = os.path.getsize(output_verification) / 1024
            logger.info("✓ Verification document created: %s (%.2f KB)",
                       output_verification, verification_size_kb)
        else:
            logger.warning("! Verification document NOT created (expected: %s)",
                         output_verification)

        # Count paragraphs in each document
        logger.info("")
        logger.info("Analyzing output documents...")

        from docx import Document

        clean_doc = Document(output_clean)
        clean_para_count = len([p for p in clean_doc.paragraphs if p.text.strip()])
        logger.info("Clean document paragraphs: %d", clean_para_count)

        if os.path.exists(output_verification):
            verification_doc = Document(output_verification)
            verification_para_count = len([p for p in verification_doc.paragraphs if p.text.strip()])
            logger.info("Verification document paragraphs: %d", verification_para_count)
            logger.info("Filtered out: %d empty/invalid paragraphs",
                       verification_para_count - clean_para_count)

        # Show sample of clean content
        logger.info("")
        logger.info("Sample of clean document content:")
        logger.info("-" * 80)
        for i, para in enumerate(clean_doc.paragraphs[:5], 1):
            if para.text.strip():
                logger.info("Paragraph %d: %s...", i, para.text[:80])
        logger.info("-" * 80)

        logger.info("")
        logger.info("=" * 80)
        logger.info("TEST COMPLETED SUCCESSFULLY")
        logger.info("=" * 80)
        logger.info("")
        logger.info("Next steps:")
        logger.info("1. Review clean document: %s", output_clean)
        logger.info("2. Compare with verification document: %s", output_verification)
        logger.info("3. Verify no excessive whitespace in clean document")
        logger.info("4. Confirm paragraph boundaries are preserved")

        return 0

    except Exception as e:
        logger.error("Processing failed: %s", e, exc_info=True)
        return 1

if __name__ == '__main__':
    exit(main())
