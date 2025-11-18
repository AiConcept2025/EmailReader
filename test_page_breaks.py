"""
Test script to verify page breaks are being added to DOCX output.
"""
import json
import os
from src.models.formatted_document import FormattedDocument
from src.convert_to_docx import convert_structured_to_docx

# Load the Konnova JSON file
print("Loading Konnova JSON file...")
with open('completed_temp/Konnova_landing_ai_20251116_192825.json', 'r') as f:
    data = json.load(f)

# Create FormattedDocument
print("Creating FormattedDocument...")
doc = FormattedDocument.from_landing_ai_response(data)

print(f"\nFormattedDocument Analysis:")
print(f"  Total pages: {doc.total_pages}")
print(f"  Total paragraphs: {doc.total_paragraphs}")

for page in doc.pages:
    print(f"\n  Page {page.page_number}:")
    print(f"    Paragraphs: {len(page.paragraphs)}")
    print(f"    Columns: {page.columns}")

    # Show font size distribution
    font_sizes = [p.font_size for p in page.paragraphs if p.font_size]
    if font_sizes:
        print(f"    Font sizes: {min(font_sizes):.1f}pt - {max(font_sizes):.1f}pt")

    # Show text types
    text_types = {}
    for para in page.paragraphs:
        text_type = para.text_type or 'unknown'
        text_types[text_type] = text_types.get(text_type, 0) + 1
    print(f"    Text types: {text_types}")

# Convert to DOCX
output_path = 'test_output_konnova.docx'
print(f"\nConverting to DOCX: {output_path}")
convert_structured_to_docx(doc, output_path)

if os.path.exists(output_path):
    file_size = os.path.getsize(output_path) / 1024
    print(f"\nDOCX created successfully: {file_size:.2f} KB")

    # Verify page breaks using python-docx
    from docx import Document
    docx_doc = Document(output_path)

    # Count page breaks
    page_break_count = 0
    for element in docx_doc.element.body:
        if element.tag.endswith('br'):
            if element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                page_break_count += 1

    # Alternative: count sections and paragraphs
    sections = len(docx_doc.sections)
    paragraphs = len(docx_doc.paragraphs)

    print(f"\nDOCX structure:")
    print(f"  Sections: {sections}")
    print(f"  Paragraphs: {paragraphs}")
    print(f"  Page breaks (explicit): {page_break_count}")

    # Check for page break runs
    page_break_runs = 0
    for para in docx_doc.paragraphs:
        for run in para.runs:
            if run._element.xpath('.//w:br[@w:type="page"]'):
                page_break_runs += 1

    print(f"  Page breaks (in runs): {page_break_runs}")

    # Show sample paragraphs from different sections
    print("\nSample paragraphs (first 3):")
    for i, para in enumerate(docx_doc.paragraphs[:3]):
        text_preview = para.text[:60].replace('\n', ' ')
        font_size = None
        if para.runs:
            font_size = para.runs[0].font.size
        print(f"  Para {i}: size={font_size}, text='{text_preview}...'")

else:
    print(f"ERROR: DOCX file was not created!")

print("\nTest complete!")
