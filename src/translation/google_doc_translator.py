"""
Google Document Translator

Uses Google Cloud Translation API v3 for document translation with format preservation.
"""

import os
import time
import logging
from typing import Dict, Any, List

from google.cloud import translate_v3
from google.api_core.exceptions import GoogleAPIError
from google.oauth2 import service_account

from src.translation.base_translator import BaseTranslator

logger = logging.getLogger('EmailReader.Translation.GoogleDoc')


class GoogleDocTranslator(BaseTranslator):
    """
    Google Cloud Translation API v3 Document Translator.

    Uses the advanced document translation endpoint that preserves
    formatting and layout in translated documents.
    """

    def __init__(self, config: Dict[str, Any]):
        """
        Initialize the Google Document translator.

        Args:
            config: Configuration dictionary with 'project_id', 'service_account', and optional 'location' or 'endpoint'

        Raises:
            ValueError: If required config values are missing
        """
        super().__init__(config)

        self.project_id = config.get('project_id')
        self.location = config.get('location', 'global')
        self.endpoint = config.get('endpoint')
        self.service_account_info = config.get('service_account')

        logger.info("=" * 80)
        logger.info("INITIALIZING GOOGLE CLOUD TRANSLATION API v3 CLIENT")
        logger.info("=" * 80)

        # Log full configuration (without sensitive data)
        logger.debug("Full configuration received:")
        logger.debug("  - project_id: %s", self.project_id)
        logger.debug("  - location: %s", self.location)
        logger.debug("  - endpoint: %s", self.endpoint)
        logger.debug("  - service_account provided: %s", bool(self.service_account_info))
        logger.debug("  - mime_type: %s", config.get('mime_type', 'N/A'))

        if self.service_account_info:
            logger.debug("  - Service account email: %s",
                        self.service_account_info.get('client_email', 'N/A'))
            logger.debug("  - Service account project: %s",
                        self.service_account_info.get('project_id', 'N/A'))

        if not self.project_id:
            raise ValueError(
                "Google Document Translation requires 'project_id' in configuration"
            )

        if not self.service_account_info:
            raise ValueError(
                "Google Document Translation requires 'service_account' credentials in configuration"
            )

        try:
            # Create credentials from service account info
            logger.info("Creating credentials from service account configuration")
            credentials = service_account.Credentials.from_service_account_info(
                self.service_account_info
            )
            logger.debug("Credentials created successfully")
            logger.debug("  - Service account: %s", credentials.service_account_email)
            logger.debug("  - Project ID: %s", credentials.project_id)

            # Prepare client options
            client_options = {}
            if self.endpoint:
                # Use custom endpoint if provided
                full_endpoint = self.endpoint
                if not full_endpoint.startswith('http'):
                    full_endpoint = f"https://{full_endpoint}"

                logger.info("Using CUSTOM endpoint configuration:")
                logger.info("  - Configured endpoint: %s", self.endpoint)
                logger.info("  - Full API endpoint URL: %s", full_endpoint)
                logger.debug("  - Expected API base: translate.googleapis.com")
                logger.debug("  - Expected regional endpoints: {region}-translate.googleapis.com")

                client_options["api_endpoint"] = self.endpoint
            else:
                logger.info("Using DEFAULT Google Cloud endpoint")
                logger.debug("  - Will use standard translate.googleapis.com")
                logger.debug("  - Location will determine regional routing")

            # Initialize client with explicit credentials
            logger.info("Initializing Translation Service Client with explicit credentials")
            if client_options:
                self.client = translate_v3.TranslationServiceClient(
                    credentials=credentials,
                    client_options=client_options
                )
            else:
                self.client = translate_v3.TranslationServiceClient(
                    credentials=credentials
                )
            logger.debug("Client created successfully")

            # Construct parent path
            self.parent = f"projects/{self.project_id}/locations/{self.location}"

            logger.info("-" * 80)
            logger.info("CLIENT INITIALIZATION SUCCESSFUL")
            logger.info("-" * 80)
            logger.info("API Configuration:")
            logger.info("  - Parent path: %s", self.parent)
            logger.info("  - Project ID: %s", self.project_id)
            logger.info("  - Location: %s", self.location)
            logger.info("  - Endpoint: %s", self.endpoint or "default (translate.googleapis.com)")
            logger.info("  - Using service account: %s", credentials.service_account_email)

            # Try to get transport info if available
            try:
                transport = getattr(self.client, '_transport', None)
                if transport:
                    logger.debug("Transport information:")
                    logger.debug("  - Transport type: %s", type(transport).__name__)
                    if hasattr(transport, '_host'):
                        logger.debug("  - Transport host: %s", transport._host)
            except Exception as transport_err:
                logger.debug("Could not retrieve transport info: %s", transport_err)

            logger.info("=" * 80)

        except Exception as e:
            logger.error("=" * 80)
            logger.error("FAILED TO INITIALIZE GOOGLE TRANSLATION CLIENT")
            logger.error("=" * 80)
            logger.error("Error type: %s", type(e).__name__)
            logger.error("Error message: %s", str(e))
            logger.error("Configuration attempted:")
            logger.error("  - project_id: %s", self.project_id)
            logger.error("  - location: %s", self.location)
            logger.error("  - endpoint: %s", self.endpoint)
            logger.error("  - service_account provided: %s", bool(self.service_account_info))
            logger.error("=" * 80)
            raise

    def translate_document(
        self,
        input_path: str,
        output_path: str,
        source_lang: str | None = None,
        target_lang: str = 'en'
    ) -> None:
        """
        Translate a DOCX document using Google Cloud Translation API v3.

        Args:
            input_path: Path to input DOCX document
            output_path: Path to save translated DOCX document
            source_lang: Source language code (e.g., 'ru', 'es') or None for auto-detect
            target_lang: Target language code (default: 'en')

        Raises:
            FileNotFoundError: If input file doesn't exist
            RuntimeError: If translation fails
        """
        logger.info("Translating document with Google Translation API v3: %s -> %s",
                   os.path.basename(input_path), target_lang)
        logger.debug("Input: %s", input_path)
        logger.debug("Output: %s", output_path)
        logger.debug("Target language: %s", target_lang)

        # Log source language handling
        if source_lang:
            logger.info("Source language: %s (specified)", source_lang)
        else:
            logger.info("Source language: auto-detect")

        if not os.path.exists(input_path):
            logger.error("Input file not found: %s", input_path)
            raise FileNotFoundError(f"File not found: {input_path}")

        # Validate file type
        if not input_path.lower().endswith('.docx'):
            logger.error("Input file must be DOCX format: %s", input_path)
            raise ValueError("Only DOCX files are supported for document translation")

        try:
            # Read document
            logger.info("Reading input document")
            with open(input_path, 'rb') as f:
                document_content = f.read()

            file_size_kb = len(document_content) / 1024
            logger.debug("Document size: %.2f KB", file_size_kb)

            # Translate
            translated_content = self._call_translation_api(document_content, target_lang, source_lang)

            # Save result
            logger.info("Saving translated document")
            with open(output_path, 'wb') as f:
                f.write(translated_content)

            if os.path.exists(output_path):
                output_size = os.path.getsize(output_path) / 1024
                logger.info("Translation completed successfully: %s (%.2f KB)",
                           os.path.basename(output_path), output_size)
            else:
                raise RuntimeError("Translation failed to create output file")

        except FileNotFoundError:
            raise
        except ValueError:
            raise
        except Exception as e:
            logger.error("Error during Google document translation: %s", e, exc_info=True)
            raise RuntimeError(f"Google document translation failed: {e}")

    def translate_document_paragraphs(self, input_path: str, output_path: str,
                                       source_lang: str = None, target_lang: str = 'en') -> str:
        """
        Translate a DOCX document using paragraph-based translation.

        This method:
        1. Extracts paragraphs from input DOCX
        2. Translates paragraphs using translate_paragraphs()
        3. Creates new DOCX with translated paragraphs

        Args:
            input_path: Path to input .docx file
            output_path: Path to save translated .docx file
            source_lang: Source language code (e.g., 'ru')
            target_lang: Target language code (default: 'en')

        Returns:
            Path to output file

        Raises:
            FileNotFoundError: If input file doesn't exist
            RuntimeError: If translation fails
        """
        from docx import Document
        from docx.shared import Pt

        logger.info("Translating document using paragraph-based approach: %s -> %s",
                   os.path.basename(input_path), target_lang)
        logger.debug("Input: %s", input_path)
        logger.debug("Output: %s", output_path)

        if not os.path.exists(input_path):
            logger.error("Input file not found: %s", input_path)
            raise FileNotFoundError(f"File not found: {input_path}")

        if not input_path.lower().endswith('.docx'):
            logger.error("Input file must be DOCX format: %s", input_path)
            raise ValueError("Only DOCX files are supported for paragraph-based translation")

        try:
            # Extract paragraphs
            logger.info("Extracting paragraphs from input document")
            doc = Document(input_path)
            paragraphs = [p.text for p in doc.paragraphs]

            logger.info("Extracted %d paragraphs from document", len(paragraphs))

            # Translate using paragraph-based API
            logger.info("Translating paragraphs using translate_paragraphs() method")
            translated_paragraphs = self.translate_paragraphs(
                paragraphs=paragraphs,
                source_lang=source_lang,
                target_lang=target_lang
            )

            logger.info("Translation complete, creating output document")

            # Create output DOCX
            output_doc = Document()
            paragraph_count = 0

            for para_text in translated_paragraphs:
                if para_text.strip():
                    p = output_doc.add_paragraph(para_text)
                    # Apply consistent formatting
                    for run in p.runs:
                        run.font.size = Pt(11)
                    # Set spacing (similar to OCR output)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    paragraph_count += 1

            logger.info("Created output document with %d paragraphs", paragraph_count)

            # Save output document
            output_doc.save(output_path)

            if os.path.exists(output_path):
                output_size = os.path.getsize(output_path) / 1024
                logger.info("Paragraph-based translation completed successfully: %s (%.2f KB)",
                           os.path.basename(output_path), output_size)
            else:
                raise RuntimeError("Translation failed to create output file")

            return output_path

        except Exception as e:
            logger.error("Error during paragraph-based document translation: %s", e, exc_info=True)
            raise RuntimeError(f"Paragraph-based document translation failed: {e}")

    def translate_paragraphs(
        self,
        paragraphs: List[str],
        target_lang: str,
        source_lang: str | None = None,
        batch_size: int = 15,
        preserve_formatting: bool = True
    ) -> List[str]:
        """
        Translate a list of paragraphs in batches using Google Cloud Translation API v3.

        This method efficiently translates multiple paragraphs by batching them together
        to minimize API calls while preserving paragraph boundaries and order.

        Args:
            paragraphs: List of paragraph text strings to translate
            target_lang: Target language code (e.g., 'en', 'es', 'fr')
            source_lang: Source language code (e.g., 'ru', 'es') or None for auto-detect
            batch_size: Number of paragraphs to process per API call (default: 15)
            preserve_formatting: Maintain paragraph boundaries (default: True)

        Returns:
            List of translated paragraph strings in the same order as input

        Raises:
            ValueError: If paragraphs is empty or target_lang is invalid
            RuntimeError: If translation fails after retries or paragraph count mismatch

        Example:
            >>> translator = GoogleDocTranslator(config)
            >>> paragraphs = ["Hello world", "How are you?"]
            >>> translated = translator.translate_paragraphs(paragraphs, "es")
            >>> print(translated)
            ['Hola mundo', '¿Cómo estás?']
        """
        if not paragraphs:
            logger.warning("Empty paragraphs list provided to translate_paragraphs")
            return []

        if not target_lang:
            raise ValueError("Target language code is required")

        input_count = len(paragraphs)

        # MANDATORY logging before translation
        logger.info("PARAGRAPH_COUNT_VERIFICATION: stage=TRANSLATION_INPUT, count=%d", input_count)
        logger.info("Starting translation: %d paragraphs", input_count)
        logger.info("Translating %d paragraphs to '%s' in batches of %d",
                   input_count, target_lang, batch_size)

        # Filter out empty paragraphs but remember their positions
        non_empty_indices = []
        non_empty_paragraphs = []

        for i, paragraph in enumerate(paragraphs):
            if paragraph and paragraph.strip():
                non_empty_indices.append(i)
                non_empty_paragraphs.append(paragraph)
            else:
                logger.debug("Skipping empty paragraph at index %d", i)

        if not non_empty_paragraphs:
            logger.warning("All paragraphs are empty, returning original list")
            return paragraphs

        logger.info("Processing %d non-empty paragraphs out of %d total",
                   len(non_empty_paragraphs), len(paragraphs))

        # Translate in batches
        translated_paragraphs = []
        total_batches = (len(non_empty_paragraphs) + batch_size - 1) // batch_size

        for batch_idx in range(0, len(non_empty_paragraphs), batch_size):
            batch_num = (batch_idx // batch_size) + 1
            batch = non_empty_paragraphs[batch_idx:batch_idx + batch_size]

            logger.info("Processing batch %d/%d (%d paragraphs)",
                       batch_num, total_batches, len(batch))

            try:
                # Translate the batch
                batch_result = self._translate_text_batch(
                    batch,
                    target_lang,
                    source_lang,
                    preserve_formatting=preserve_formatting
                )
                translated_paragraphs.extend(batch_result)

                logger.debug("Batch %d/%d completed successfully", batch_num, total_batches)

            except Exception as e:
                logger.warning("Batch %d/%d failed: %s. Retrying with individual translations.",
                              batch_num, total_batches, str(e))

                # Fallback: translate paragraphs individually
                for para in batch:
                    try:
                        individual_result = self._translate_text_batch(
                            [para],
                            target_lang,
                            source_lang,
                            preserve_formatting=preserve_formatting
                        )
                        translated_paragraphs.extend(individual_result)
                    except Exception as individual_error:
                        logger.error("Failed to translate paragraph: %s", str(individual_error))
                        logger.warning("Using original text as fallback for failed paragraph")
                        # Keep original paragraph if translation fails (NEVER skip)
                        translated_paragraphs.append(para)

        # Reconstruct full paragraph list with empty paragraphs in original positions
        result = list(paragraphs)  # Start with original list (preserves empty strings)

        for i, translated in zip(non_empty_indices, translated_paragraphs):
            result[i] = translated

        output_count = len(result)

        # MANDATORY logging after translation
        logger.info("PARAGRAPH_COUNT_VERIFICATION: stage=TRANSLATION_OUTPUT, count=%d", output_count)
        logger.info("Translation complete: %d paragraphs received, %d paragraphs translated",
                   input_count, output_count)

        # RAISE EXCEPTION if counts don't match
        if input_count != output_count:
            error_msg = (
                f"Paragraph count mismatch in translation: "
                f"input={input_count}, output={output_count}"
            )
            logger.error(error_msg)
            raise RuntimeError(error_msg)

        return result

    def _translate_text_batch(
        self,
        texts: List[str],
        target_lang: str,
        source_lang: str | None = None,
        preserve_formatting: bool = True
    ) -> List[str]:
        """
        Translate a batch of text strings using Google Cloud Translation API v3.

        Internal method that handles the actual API call for text translation.
        Uses a delimiter-based approach to maintain paragraph boundaries.

        Args:
            texts: List of text strings to translate
            target_lang: Target language code
            source_lang: Source language code or None for auto-detect
            preserve_formatting: Maintain paragraph boundaries

        Returns:
            List of translated text strings

        Raises:
            RuntimeError: If API call fails
        """
        if not texts:
            return []

        # Use a unique delimiter to separate paragraphs
        delimiter = "\n\n###PARAGRAPH_BOUNDARY###\n\n"

        if preserve_formatting:
            # Join texts with delimiter
            combined_text = delimiter.join(texts)
            logger.debug("Combined %d texts into single string (%d chars)",
                        len(texts), len(combined_text))
        else:
            combined_text = "\n\n".join(texts)

        try:
            start_time = time.time()

            logger.debug("Calling Translation API for text batch")
            logger.debug("  - Text length: %d characters", len(combined_text))
            logger.debug("  - Target language: %s", target_lang)
            logger.debug("  - Preserve formatting: %s", preserve_formatting)

            # Log source language handling
            if source_lang:
                logger.debug("Using specified source language: %s", source_lang)
            else:
                logger.debug("Source language not specified - API will auto-detect")

            # Call the text translation API (not document translation)
            request = translate_v3.TranslateTextRequest(
                parent=self.parent,
                contents=[combined_text],
                source_language_code=source_lang if source_lang else None,  # Explicit None if not specified
                target_language_code=target_lang,
                mime_type="text/plain"
            )

            response = self.client.translate_text(request=request)

            duration = time.time() - start_time
            logger.debug("API call completed in %.2f seconds", duration)

            if not response.translations:
                raise RuntimeError("API returned empty translation response")

            translated_text = response.translations[0].translated_text
            logger.debug("Received translated text (%d chars)", len(translated_text))

            # Split back into paragraphs
            if preserve_formatting and delimiter in translated_text:
                result = translated_text.split(delimiter)
                logger.debug("Split translated text into %d paragraphs", len(result))
            else:
                # Fallback: split by double newlines
                result = [p.strip() for p in translated_text.split("\n\n") if p.strip()]
                logger.debug("Fallback split resulted in %d paragraphs", len(result))

            # Ensure we have the same number of results as inputs
            if len(result) != len(texts):
                logger.warning(
                    "Translation result count mismatch: expected %d, got %d. "
                    "Falling back to individual translation.",
                    len(texts), len(result)
                )
                # Fallback: translate each text individually
                result = []
                for text in texts:
                    individual_request = translate_v3.TranslateTextRequest(
                        parent=self.parent,
                        contents=[text],
                        source_language_code=source_lang,
                        target_language_code=target_lang,
                        mime_type="text/plain"
                    )
                    individual_response = self.client.translate_text(request=individual_request)
                    result.append(individual_response.translations[0].translated_text)

            return result

        except GoogleAPIError as e:
            logger.error("Google Translation API error in text batch: %s", str(e))
            logger.error("  - Status code: %s",
                        e.grpc_status_code if hasattr(e, 'grpc_status_code') else 'unknown')
            raise RuntimeError(f"Text batch translation failed: {e}")

        except Exception as e:
            logger.error("Unexpected error in text batch translation: %s", str(e), exc_info=True)
            raise RuntimeError(f"Text batch translation failed: {e}")

    def _call_translation_api(
        self,
        content: bytes,
        target_lang: str,
        source_lang: str = None
    ) -> bytes:
        """
        Call Google Cloud Translation API for document translation.

        Args:
            content: Document content as bytes
            target_lang: Target language code
            source_lang: Optional source language code (auto-detect if None)

        Returns:
            Translated document content as bytes

        Raises:
            RuntimeError: If API call fails
        """
        logger.info("=" * 80)
        logger.info("CALLING GOOGLE TRANSLATION API v3")
        logger.info("=" * 80)

        # Log request details
        logger.info("Request Configuration:")
        logger.info("  - Parent path: %s", self.parent)
        logger.info("  - Target language: %s", target_lang)
        logger.info("  - Source language: %s", source_lang or "auto-detect")
        logger.info("  - Document size: %d bytes (%.2f KB)", len(content), len(content) / 1024)

        # Log API endpoint details
        logger.debug("API Endpoint Details:")
        logger.debug("  - Project ID: %s", self.project_id)
        logger.debug("  - Location: %s", self.location)
        logger.debug("  - Endpoint: %s", self.endpoint or "default")

        # Construct the full API URL that will be called
        if self.endpoint:
            base_url = self.endpoint if self.endpoint.startswith('http') else f"https://{self.endpoint}"
        else:
            base_url = "https://translate.googleapis.com"

        api_url = f"{base_url}/v3/{self.parent}:translateDocument"
        logger.debug("  - Full API URL: %s", api_url)

        try:
            start_time = time.time()

            # Prepare document input
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            logger.debug("Request Payload Configuration:")
            logger.debug("  - MIME type: %s", mime_type)
            logger.debug("  - Content length: %d bytes", len(content))
            logger.debug("  - Content type: bytes")

            document_input_config = translate_v3.DocumentInputConfig(
                content=content,
                mime_type=mime_type
            )

            logger.info("Using Google Cloud default translation model (Basic NMT)")

            # Prepare request (without model parameter - uses Google's default Basic NMT)
            request = translate_v3.TranslateDocumentRequest(
                parent=self.parent,
                target_language_code=target_lang,
                document_input_config=document_input_config
            )

            if source_lang:
                request.source_language_code = source_lang
                logger.info("Using specified source language: %s", source_lang)
            else:
                logger.info("Source language not specified - API will auto-detect")

            logger.info("-" * 80)
            logger.info("SENDING REQUEST TO GOOGLE TRANSLATION API")
            logger.info("-" * 80)
            logger.debug("Request object:")
            logger.debug("  - Type: %s", type(request).__name__)
            logger.debug("  - Parent: %s", request.parent)
            logger.debug("  - Target language: %s", request.target_language_code)
            logger.debug("  - Source language: %s", getattr(request, 'source_language_code', 'auto-detect'))

            # Log transport/connection info
            try:
                transport = getattr(self.client, '_transport', None)
                if transport and hasattr(transport, '_host'):
                    logger.debug("Connection details:")
                    logger.debug("  - Transport host: %s", transport._host)
                    logger.debug("  - Transport type: %s", type(transport).__name__)
            except Exception as conn_err:
                logger.debug("Could not retrieve connection details: %s", conn_err)

            # Call API
            logger.info("Making API call to translate_document()...")
            response = self.client.translate_document(request=request)

            duration = time.time() - start_time

            logger.info("-" * 80)
            logger.info("API CALL SUCCESSFUL")
            logger.info("-" * 80)
            logger.info("Response received in %.2f seconds", duration)

            # Extract translated document
            translated_content = response.document_translation.byte_stream_outputs[0]

            logger.info("Translation Results:")
            logger.info("  - Translated content size: %d bytes (%.2f KB)",
                       len(translated_content), len(translated_content) / 1024)
            logger.info("  - Size change: %+d bytes (%.1f%%)",
                       len(translated_content) - len(content),
                       ((len(translated_content) - len(content)) / len(content)) * 100)

            # Log detected source language if auto-detected
            if not source_lang and hasattr(response, 'document_translation'):
                detected_lang = getattr(
                    response.document_translation,
                    'detected_language_code',
                    'unknown'
                )
                logger.info("  - Detected source language: %s", detected_lang)

            # Log MIME type if available
            if hasattr(response.document_translation, 'mime_type'):
                logger.debug("  - Response MIME type: %s", response.document_translation.mime_type)

            logger.info("=" * 80)

            return translated_content

        except GoogleAPIError as e:
            logger.error("=" * 80)
            logger.error("GOOGLE TRANSLATION API ERROR")
            logger.error("=" * 80)
            logger.error("Error Type: GoogleAPIError")
            logger.error("Error Details:")
            logger.error("  - Status code: %s", e.grpc_status_code if hasattr(e, 'grpc_status_code') else 'unknown')
            logger.error("  - Message: %s", e.message if hasattr(e, 'message') else str(e))
            logger.error("  - Full error: %s", str(e))

            # Log request details for debugging
            logger.error("Request Details (for debugging):")
            logger.error("  - API URL: %s", api_url)
            logger.error("  - Parent: %s", self.parent)
            logger.error("  - Project ID: %s", self.project_id)
            logger.error("  - Location: %s", self.location)
            logger.error("  - Target language: %s", target_lang)
            logger.error("  - Endpoint: %s", self.endpoint or "default")

            # Special handling for Translation LLM error
            error_str = str(e)
            if "LLM models is not enabled" in error_str or "generalModels.docPredict" in error_str:
                logger.error("")
                logger.error("TRANSLATION LLM NOT ENABLED")
                logger.error("=" * 80)
                logger.error("To use Translation LLM with document translation:")
                logger.error("1. Enable Google Cloud Advanced Translation in Cloud Console")
                logger.error("2. Navigate to: https://console.cloud.google.com/apis/library/translate.googleapis.com")
                logger.error("3. Ensure 'Advanced' tier is enabled (not just 'Basic')")
                logger.error("4. Check billing is enabled and quota is sufficient")
                logger.error("")
                logger.error("Alternative: Change config to use NMT model:")
                logger.error("  'model': 'nmt'  (or remove model parameter)")
                logger.error("=" * 80)

            # Check if it's an authentication error
            if hasattr(e, 'grpc_status_code'):
                from grpc import StatusCode
                if e.grpc_status_code == StatusCode.UNAUTHENTICATED:
                    logger.error("AUTHENTICATION ERROR - Check service account credentials in config")
                    logger.error("  - Service account email: %s",
                               self.service_account_info.get('client_email', 'NOT SET') if self.service_account_info else 'NOT PROVIDED')
                    logger.error("  - Service account project: %s",
                               self.service_account_info.get('project_id', 'NOT SET') if self.service_account_info else 'NOT PROVIDED')
                elif e.grpc_status_code == StatusCode.PERMISSION_DENIED:
                    logger.error("PERMISSION DENIED - Check API permissions and quota")
                    logger.error("  - Ensure Translation API is enabled for project: %s", self.project_id)
                    logger.error("  - Verify service account has 'Cloud Translation API User' role")
                elif e.grpc_status_code == StatusCode.NOT_FOUND:
                    logger.error("NOT FOUND - Check project ID, location, and endpoint")
                elif e.grpc_status_code == StatusCode.INVALID_ARGUMENT:
                    logger.error("INVALID ARGUMENT - Check request parameters")

            logger.error("=" * 80)
            raise RuntimeError(f"Google Translation API failed: {e}")

        except Exception as e:
            logger.error("=" * 80)
            logger.error("UNEXPECTED ERROR CALLING TRANSLATION API")
            logger.error("=" * 80)
            logger.error("Error Type: %s", type(e).__name__)
            logger.error("Error Message: %s", str(e))
            logger.error("Request was:")
            logger.error("  - API URL: %s", api_url)
            logger.error("  - Parent: %s", self.parent)
            logger.error("  - Target language: %s", target_lang)
            logger.error("=" * 80)
            logger.error("Stack trace:", exc_info=True)
            raise RuntimeError(f"Translation API call failed: {e}")
