"""
Google Batch Translator

Uses Google Cloud Translation API v3 translate_text() endpoint for paragraph-based
batch translation with format preservation.
"""

import os
import time
import logging
from typing import Dict, Any, List

from google.cloud import translate_v3
from google.api_core.exceptions import GoogleAPIError
from google.oauth2 import service_account
from docx import Document
from docx.shared import Pt

from src.translation.base_translator import BaseTranslator
from src.models.paragraph import Paragraph, TextSpan
from src.config import load_config
from src.translation.claude_validator.validator import ClaudeTranslationValidator

logger = logging.getLogger('EmailReader.Translation.GoogleBatch')


class GoogleBatchTranslator(BaseTranslator):
    """
    Google Cloud Translation API v3 Batch Translator.

    Translates documents paragraph-by-paragraph using the translate_text() API
    in batches for efficient processing while preserving formatting.
    """

    BATCH_SIZE = 25  # Number of paragraphs to translate per API call

    def __init__(self, config: Dict[str, Any]):
        """
        Initialize the Google Batch translator.

        Args:
            config: Configuration dictionary with 'project_id' and optional 'location'

        Raises:
            ValueError: If required config values are missing
        """
        super().__init__(config)

        self.project_id = config.get('project_id')
        self.location = config.get('location', 'us-central1')

        if not self.project_id:
            raise ValueError(
                "Google Batch Translation requires 'project_id' in configuration"
            )

        logger.info("Initializing Google Cloud Translation API v3 Batch client")
        logger.debug("Project ID: %s", self.project_id)
        logger.debug("Location: %s", self.location)
        logger.debug("Batch size: %d paragraphs", self.BATCH_SIZE)

        try:
            # Load service account credentials directly from config.dev.json
            full_config = load_config()
            sa_info = full_config.get('google_drive', {}).get('service_account')

            if not sa_info:
                raise ValueError(
                    "Service account credentials not found in config. "
                    "Please ensure 'google_drive.service_account' exists in config.dev.json"
                )

            # Create credentials directly from config dict (no temp file needed)
            credentials = service_account.Credentials.from_service_account_info(sa_info)
            logger.debug("Using service account credentials from config.dev.json")

            # Create client with credentials
            self.client = translate_v3.TranslationServiceClient(credentials=credentials)
            self.parent = f"projects/{self.project_id}/locations/{self.location}"
            logger.info("Google Batch Translator initialized successfully")

            # NEW: Initialize Claude validator if enabled
            validation_config = full_config.get('translation', {}).get('validation', {})
            self.validation_enabled = validation_config.get('enabled', False)
            self.validator = None

            if self.validation_enabled:
                config_path = validation_config.get('config_path')
                if config_path:
                    try:
                        self.validator = ClaudeTranslationValidator(config_path)
                        logger.info("Claude post-translation validation enabled")
                    except Exception as validator_error:
                        logger.warning(
                            "Failed to initialize Claude validator: %s. "
                            "Translation will proceed without validation.",
                            validator_error
                        )
                        self.validation_enabled = False
                else:
                    logger.warning("Claude validation enabled but config_path not provided")
                    self.validation_enabled = False

        except Exception as e:
            logger.error("Failed to initialize Google Translation client: %s", e)
            raise

    def translate_document(
        self,
        input_path: str,
        output_path: str,
        target_lang: str = 'en'
    ) -> None:
        """
        Translate a DOCX document using paragraph-based batch translation.

        Args:
            input_path: Path to input DOCX document
            output_path: Path to save translated DOCX document
            target_lang: Target language code (default: 'en')

        Raises:
            FileNotFoundError: If input file doesn't exist
            RuntimeError: If translation fails
        """
        logger.info("[GoogleBatchTranslator] STARTING TRANSLATION: %s -> %s",
                   os.path.basename(input_path), target_lang)
        logger.debug("Input: %s", input_path)
        logger.debug("Output: %s", output_path)
        logger.debug("Target language: %s", target_lang)

        if not os.path.exists(input_path):
            logger.error("Input file not found: %s", input_path)
            raise FileNotFoundError(f"File not found: {input_path}")

        input_size_kb = os.path.getsize(input_path) / 1024
        logger.info("[GoogleBatchTranslator] Input file size: %.2f KB", input_size_kb)

        try:
            # Extract paragraphs from DOCX
            logger.info("[GoogleBatchTranslator] Extracting paragraphs from DOCX")
            paragraphs = self._extract_paragraphs_from_docx(input_path)
            logger.info("[GoogleBatchTranslator] Extracted %d paragraphs from input", len(paragraphs))

            # Log paragraph content for debugging
            if paragraphs:
                sample_para = paragraphs[0]
                logger.info("[GoogleBatchTranslator] Sample paragraph (first): role=%s, content_len=%d",
                           sample_para.role, len(sample_para.content))

            # Translate paragraphs in batches
            total_batches = (len(paragraphs) + self.BATCH_SIZE - 1) // self.BATCH_SIZE
            logger.info("[GoogleBatchTranslator] Processing %d batches of %d paragraphs each",
                       total_batches, self.BATCH_SIZE)

            translated_paragraphs = []
            batch_start_time = time.time()

            for batch_idx in range(total_batches):
                start_idx = batch_idx * self.BATCH_SIZE
                end_idx = min(start_idx + self.BATCH_SIZE, len(paragraphs))
                batch = paragraphs[start_idx:end_idx]

                logger.info("[GoogleBatchTranslator] Processing batch %d/%d (%d paragraphs)",
                           batch_idx + 1, total_batches, len(batch))

                # Extract text content from paragraphs
                texts = [p.content for p in batch]

                # Translate batch
                batch_translate_start = time.time()
                translated_texts = self._translate_batch(texts, target_lang)
                batch_duration = time.time() - batch_translate_start

                logger.info("[GoogleBatchTranslator] Batch %d translated in %.2f seconds",
                           batch_idx + 1, batch_duration)

                # Map translated text back to Paragraph objects
                for para, translated_text in zip(batch, translated_texts):
                    translated_para = Paragraph(
                        content=translated_text,
                        page=para.page,
                        role=para.role,
                        spans=para.spans,  # Preserve original formatting spans
                        bounding_box=para.bounding_box,
                        is_list_item=para.is_list_item,
                        list_marker=para.list_marker
                    )
                    translated_paragraphs.append(translated_para)

            total_duration = time.time() - batch_start_time
            logger.info("[GoogleBatchTranslator] All batches completed in %.2f seconds",
                       total_duration)

            # Log translation verification
            if translated_paragraphs:
                logger.info("[GoogleBatchTranslator] Translated %d paragraphs successfully",
                           len(translated_paragraphs))
                sample_translated = translated_paragraphs[0]
                logger.info("[GoogleBatchTranslator] Sample translated paragraph (first): role=%s, content_len=%d",
                           sample_translated.role, len(sample_translated.content))
            else:
                logger.warning("[GoogleBatchTranslator] WARNING: No translated paragraphs!")

            # Filter problematic characters from translated text (BEFORE validation)
            # This ensures Claude validator receives clean text without formatting-breaking characters
            self._filter_problematic_characters(translated_paragraphs)

            # NEW: Claude post-translation validation (quality check, NOT re-translation)
            if self.validation_enabled and self.validator and translated_paragraphs:
                logger.info("[GoogleBatchTranslator] Starting Claude post-translation validation...")

                try:
                    validation_start_time = time.time()
                    validation_result = self.validator.validate_translations(
                        translated_paragraphs=translated_paragraphs,
                        target_lang=target_lang
                    )
                    validation_duration = time.time() - validation_start_time

                    # Replace translated content with validated versions (grammar/format fixes)
                    for i, validated_text in enumerate(validation_result.validated_texts):
                        if i < len(translated_paragraphs):
                            translated_paragraphs[i].content = validated_text

                    # Log validation metrics
                    logger.info(
                        "[GoogleBatchTranslator] Validation complete in %.2f seconds: %s",
                        validation_duration,
                        validation_result.quality_metrics
                    )
                    logger.info(
                        "[GoogleBatchTranslator] Cache stats: %s",
                        validation_result.cache_stats
                    )

                except Exception as validation_error:
                    logger.warning(
                        "[GoogleBatchTranslator] Validation failed: %s. "
                        "Proceeding with Google-translated text.",
                        validation_error
                    )

            # Reassemble DOCX
            logger.info("[GoogleBatchTranslator] Assembling translated DOCX from %d paragraphs",
                       len(translated_paragraphs))
            self._assemble_docx(translated_paragraphs, output_path)

            # Verify output file was created
            try:
                if os.path.exists(output_path):
                    output_size = os.path.getsize(output_path) / 1024
                    logger.info("[GoogleBatchTranslator] Translation completed: %s (%.2f KB)",
                               os.path.basename(output_path), output_size)
                else:
                    logger.warning("Output file check: %s does not exist", output_path)
            except OSError as e:
                # In test environments, file might not exist - that's ok
                logger.debug("Could not verify output file: %s", e)

        except FileNotFoundError:
            raise
        except Exception as e:
            logger.error("Error during Google batch translation: %s", e, exc_info=True)
            raise RuntimeError(f"Google batch translation failed: {e}")

    def _translate_batch(self, texts: List[str], target_lang: str) -> List[str]:
        """
        Translate a batch of text strings using Google Translation API.

        Handles empty strings by skipping them in the API call and preserving
        their positions in the returned list.

        Args:
            texts: List of text strings to translate (may include empty strings)
            target_lang: Target language code

        Returns:
            List of translated text strings (with empty strings preserved in original positions)

        Raises:
            RuntimeError: If API call fails
        """
        logger.debug("[GoogleBatchTranslator] Translating batch of %d texts", len(texts))
        logger.debug("Parent: %s", self.parent)
        logger.debug("Target language: %s", target_lang)

        # Filter out empty strings and track their positions
        non_empty_indices = [i for i, text in enumerate(texts) if text.strip()]
        non_empty_texts = [texts[i] for i in non_empty_indices]

        empty_count = len(texts) - len(non_empty_texts)
        if empty_count > 0:
            logger.debug("[GoogleBatchTranslator] Skipping %d empty texts (preserving for spacing)", empty_count)

        # If all texts are empty, return them as-is
        if not non_empty_texts:
            logger.debug("[GoogleBatchTranslator] All texts empty, returning without API call")
            return texts

        try:
            start_time = time.time()

            # Prepare request with only non-empty texts
            request = translate_v3.TranslateTextRequest(
                parent=self.parent,
                contents=non_empty_texts,
                target_language_code=target_lang,
                mime_type='text/plain'
            )

            logger.debug("[GoogleBatchTranslator] Sending batch translation request to Google API (%d texts)",
                        len(non_empty_texts))

            # Call API
            response = self.client.translate_text(request=request)

            duration = time.time() - start_time
            logger.debug("[GoogleBatchTranslator] Batch translation completed in %.2f seconds", duration)

            # Extract translated texts
            translated_non_empty = [t.translated_text for t in response.translations]

            logger.debug("[GoogleBatchTranslator] Received %d translated texts", len(translated_non_empty))

            # Reconstruct full list with empty strings in original positions
            translated_texts = [""] * len(texts)
            for idx, translated in zip(non_empty_indices, translated_non_empty):
                translated_texts[idx] = translated

            return translated_texts

        except Exception as e:
            # Check if it's a GoogleAPIError
            if type(e).__name__ == 'GoogleAPIError' or isinstance(e, GoogleAPIError):
                logger.error("Google Translation API error: %s", e)
                logger.error("Error details: status=%s, message=%s",
                            e.grpc_status_code if hasattr(e, 'grpc_status_code') else 'unknown',
                            e.message if hasattr(e, 'message') else str(e))
                raise RuntimeError(f"Translation API failed: {e}")
            else:
                logger.error("Unexpected error calling Translation API: %s", e, exc_info=True)
                raise RuntimeError(f"Translation API call failed: {e}")

    def _extract_paragraphs_from_docx(self, input_path: str) -> List[Paragraph]:
        """
        Extract paragraphs from a DOCX file.

        Args:
            input_path: Path to DOCX file

        Returns:
            List of Paragraph objects with formatting metadata

        Raises:
            RuntimeError: If extraction fails
        """
        logger.debug("[GoogleBatchTranslator] Extracting paragraphs from: %s", input_path)

        try:
            doc = Document(input_path)
            paragraphs = []

            for para in doc.paragraphs:
                # Map style to role
                style_name = para.style.name if para.style else 'Normal'
                role = self._map_style_to_role(style_name)

                # Handle empty paragraphs - preserve them for spacing
                if not para.text.strip():
                    # Create empty paragraph object to preserve document structure
                    paragraph = Paragraph(
                        content="",  # Keep empty for spacing
                        page=0,
                        role=role,
                        spans=[]
                    )
                    paragraphs.append(paragraph)
                    logger.debug("[GoogleBatchTranslator] Preserving empty paragraph for spacing")
                    continue

                # Extract runs as text spans for non-empty paragraphs
                spans = self._extract_runs_as_spans(para.runs)

                # Create Paragraph object
                paragraph = Paragraph(
                    content=para.text,
                    page=0,  # DOCX doesn't have page info
                    role=role,
                    spans=spans
                )

                paragraphs.append(paragraph)

            logger.debug("[GoogleBatchTranslator] Extracted %d paragraphs", len(paragraphs))
            return paragraphs

        except Exception as e:
            logger.error("Error extracting paragraphs from DOCX: %s", e, exc_info=True)
            raise RuntimeError(f"Failed to extract paragraphs from DOCX: {e}")

    def _map_style_to_role(self, style_name: str) -> str:
        """
        Map DOCX style name to paragraph role.

        Args:
            style_name: DOCX style name

        Returns:
            Paragraph role string
        """
        if style_name == 'Title':
            return 'title'
        elif style_name.startswith('Heading'):
            return 'heading'
        elif style_name == 'List Paragraph':
            return 'listItem'
        else:
            return 'paragraph'

    def _extract_runs_as_spans(self, runs) -> List[TextSpan]:
        """
        Extract formatting runs as TextSpan objects.

        Args:
            runs: List of python-docx run objects

        Returns:
            List of TextSpan objects with formatting metadata
        """
        spans = []

        for run in runs:
            # Extract formatting attributes
            is_bold = run.bold if run.bold is not None else False
            is_italic = run.italic if run.italic is not None else False
            font_size = None

            if run.font.size is not None:
                font_size = run.font.size.pt

            # Create TextSpan
            span = TextSpan(
                text=run.text,
                is_bold=is_bold,
                is_italic=is_italic,
                font_size=font_size
            )

            spans.append(span)

        return spans

    def _assemble_docx(self, paragraphs: List[Paragraph], output_path: str) -> None:
        """
        Assemble translated paragraphs into a DOCX file.

        Args:
            paragraphs: List of translated Paragraph objects
            output_path: Path to save output DOCX file

        Raises:
            RuntimeError: If assembly fails
        """
        logger.debug("[GoogleBatchTranslator] Assembling %d paragraphs into DOCX", len(paragraphs))

        try:
            doc = Document()

            for idx, para in enumerate(paragraphs, 1):
                # Handle empty paragraphs (for spacing)
                if not para.content.strip():
                    doc.add_paragraph("")
                    logger.debug("[GoogleBatchTranslator] Added empty paragraph for spacing")
                    continue

                # Handle non-empty paragraphs based on role
                if para.role == 'title':
                    # Add as title (level 0 heading)
                    doc.add_heading(para.content, level=0)
                    logger.debug("[GoogleBatchTranslator] Added title paragraph")

                elif para.role == 'heading':
                    # Add as heading (level 1)
                    doc.add_heading(para.content, level=1)
                    logger.debug("[GoogleBatchTranslator] Added heading paragraph")

                elif para.role == 'listItem':
                    # Add as list item with marker
                    list_text = f"• {para.content}"
                    doc.add_paragraph(list_text)
                    logger.debug("[GoogleBatchTranslator] Added list item paragraph")

                elif para.role == 'paragraph':
                    # Add as regular paragraph
                    # CRITICAL FIX: Always use para.content (translated text), NOT span.text (original)
                    # After translation, spans contain original language text which must be discarded
                    doc.add_paragraph(para.content)
                    logger.debug("[GoogleBatchTranslator] Added paragraph (formatting discarded after translation)")

                else:
                    # Default: add as paragraph
                    doc.add_paragraph(para.content)
                    logger.debug("[GoogleBatchTranslator] Added default paragraph")

            # Save document
            doc.save(output_path)
            logger.debug("[GoogleBatchTranslator] Document saved to: %s", output_path)

        except Exception as e:
            logger.error("Error assembling DOCX: %s", e, exc_info=True)
            raise RuntimeError(f"Failed to assemble DOCX: {e}")

    def _filter_problematic_characters(self, paragraphs: List[Paragraph]) -> None:
        """
        Filter out problematic characters from paragraph content.

        These characters can break formatting when introduced by translation.
        Filters are applied in-place to paragraph content.

        Args:
            paragraphs: List of Paragraph objects to filter

        Common problematic characters:
            - '←' (U+2190): Left arrow, often appears at line ends
            - '→' (U+2192): Right arrow
            - '\u200b': Zero-width space
            - '\u00ad': Soft hyphen
        """
        # Define characters to filter out (can be extended as needed)
        problematic_chars = [
            '\u2190',  # ← Left arrow
            '\u2192',  # → Right arrow
            '\u2194',  # ↔ Left-right arrow
            '\u21d0',  # ⇐ Leftwards double arrow
            '\u21d2',  # ⇒ Rightwards double arrow
            '\u21d4',  # ⇔ Left-right double arrow
            '\u200b',  # Zero-width space
            '\u00ad',  # Soft hyphen
            '\ufeff',  # Zero-width no-break space (BOM)
        ]

        filtered_count = 0
        total_chars_removed = 0

        for para in paragraphs:
            if not para.content:
                continue

            original_content = para.content
            filtered_content = original_content

            # Remove each problematic character
            for char in problematic_chars:
                if char in filtered_content:
                    filtered_content = filtered_content.replace(char, '')

            # Update paragraph content if changes were made
            if filtered_content != original_content:
                chars_removed = len(original_content) - len(filtered_content)
                total_chars_removed += chars_removed
                filtered_count += 1
                para.content = filtered_content

                logger.debug(
                    "[GoogleBatchTranslator] Filtered %d problematic char(s) from paragraph",
                    chars_removed
                )

        if filtered_count > 0:
            logger.info(
                "[GoogleBatchTranslator] Character filtering: removed %d char(s) from %d paragraph(s)",
                total_chars_removed,
                filtered_count
            )
        else:
            logger.debug("[GoogleBatchTranslator] Character filtering: no problematic characters found")
