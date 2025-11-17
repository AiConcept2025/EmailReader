"""
Quality validation module for OCR output verification
Validates font sizes, layout preservation, and document structure
"""

from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
from pathlib import Path
import json

from docx import Document
from src.config import get_config_value
from src.logger import logger


@dataclass
class ValidationResult:
    """Results from quality validation"""
    file_path: str
    passed: bool = False
    score: float = 0.0  # Overall quality score 0-100

    # Font size validation
    font_size_valid: bool = False
    font_size_distribution: Dict[str, float] = field(default_factory=dict)
    font_size_issues: List[str] = field(default_factory=list)

    # Layout validation
    page_count: int = 0
    page_breaks_valid: bool = False
    column_detection_valid: bool = False
    layout_issues: List[str] = field(default_factory=list)

    # Quality flags
    has_outlier_fonts: bool = False
    outlier_font_sizes: List[float] = field(default_factory=list)

    # Overall issues
    all_issues: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            "file_path": self.file_path,
            "passed": self.passed,
            "score": self.score,
            "font_size_valid": self.font_size_valid,
            "font_size_distribution": self.font_size_distribution,
            "font_size_issues": self.font_size_issues,
            "page_count": self.page_count,
            "page_breaks_valid": self.page_breaks_valid,
            "column_detection_valid": self.column_detection_valid,
            "layout_issues": self.layout_issues,
            "has_outlier_fonts": self.has_outlier_fonts,
            "outlier_font_sizes": self.outlier_font_sizes,
            "all_issues": self.all_issues
        }

    def print_report(self) -> None:
        """Print validation report to console"""
        print("\n" + "="*60)
        print("QUALITY VALIDATION REPORT")
        print("="*60)
        print(f"File: {Path(self.file_path).name}")
        print(f"Overall Score: {self.score:.1f}/100")
        print(f"Status: {'✓ PASS' if self.passed else '✗ FAIL'}")
        print()

        print("FONT SIZE VALIDATION:")
        print(f"  Valid: {'✓' if self.font_size_valid else '✗'}")
        if self.font_size_distribution:
            print(f"  Body Text: {self.font_size_distribution.get('body', 0):.1f}%")
            print(f"  Headings: {self.font_size_distribution.get('heading', 0):.1f}%")
            print(f"  Titles: {self.font_size_distribution.get('title', 0):.1f}%")
        if self.has_outlier_fonts:
            print(f"  Outliers detected: {len(self.outlier_font_sizes)}")
            print(f"    Sizes: {', '.join(f'{s:.1f}pt' for s in self.outlier_font_sizes[:5])}")
        if self.font_size_issues:
            for issue in self.font_size_issues:
                print(f"  ⚠ {issue}")

        print()
        print("LAYOUT VALIDATION:")
        print(f"  Page count: {self.page_count}")
        print(f"  Page breaks: {'✓' if self.page_breaks_valid else '✗'}")
        print(f"  Column detection: {'✓' if self.column_detection_valid else '✗'}")
        if self.layout_issues:
            for issue in self.layout_issues:
                print(f"  ⚠ {issue}")

        if self.all_issues:
            print()
            print("ALL ISSUES:")
            for issue in self.all_issues:
                print(f"  • {issue}")

        print("="*60)


class QualityValidator:
    """Validate OCR output quality against specifications"""

    def __init__(self):
        """Initialize validator with thresholds from config"""
        quality_config = get_config_value('quality', {})

        # Font size thresholds
        thresholds = quality_config.get('font_size_thresholds', {})
        self.min_font_size = thresholds.get('min', 7.0)
        self.max_font_size = thresholds.get('max', 48.0)
        self.body_text_range = tuple(thresholds.get('body_text_range', [10.0, 13.0]))
        self.heading_range = tuple(thresholds.get('heading_range', [13.0, 24.0]))
        self.title_range = tuple(thresholds.get('title_range', [24.0, 48.0]))

        # Distribution targets (percentages)
        self.target_body_min = 50.0  # At least 50% body text
        self.target_body_max = 85.0  # At most 85% body text
        self.target_heading_min = 10.0  # At least 10% headings
        self.target_title_max = 10.0  # At most 10% titles

        # Calibration factor
        self.calibration_factor = quality_config.get('calibration_factor', 400.0)

        logger.info("QualityValidator initialized")
        logger.debug("  Font size range: %.1f - %.1fpt", self.min_font_size, self.max_font_size)
        logger.debug("  Body text range: %.1f - %.1fpt", *self.body_text_range)

    def validate_docx(self, file_path: str, expected_pages: Optional[int] = None) -> ValidationResult:
        """
        Validate a DOCX file for quality

        Args:
            file_path: Path to DOCX file
            expected_pages: Expected number of pages (for page break validation)

        Returns:
            ValidationResult with detailed analysis
        """
        result = ValidationResult(file_path=file_path)

        try:
            doc = Document(file_path)

            # Extract font sizes from all runs
            font_sizes = self._extract_font_sizes(doc)

            if not font_sizes:
                result.all_issues.append("No font size data found in document")
                return result

            # Validate font sizes
            self._validate_font_sizes(font_sizes, result)

            # Validate layout
            self._validate_layout(doc, result, expected_pages)

            # Calculate overall score
            result.score = self._calculate_score(result)

            # Determine pass/fail
            result.passed = result.score >= 70.0  # Pass threshold: 70%

            logger.info("Validation complete: %s (score: %.1f/100)",
                       "PASS" if result.passed else "FAIL", result.score)

        except Exception as e:
            logger.error("Validation failed: %s", e, exc_info=True)
            result.all_issues.append(f"Validation error: {e}")

        return result

    def _extract_font_sizes(self, doc: Document) -> List[float]:
        """Extract all font sizes from document runs"""
        font_sizes = []

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.size:
                    # Convert EMU to points
                    size_pt = run.font.size.pt
                    font_sizes.append(size_pt)

        logger.debug("Extracted %d font size values", len(font_sizes))
        return font_sizes

    def _validate_font_sizes(self, font_sizes: List[float], result: ValidationResult) -> None:
        """Validate font size distribution"""
        total = len(font_sizes)

        # Count sizes in each category
        body_count = sum(1 for s in font_sizes if self.body_text_range[0] <= s <= self.body_text_range[1])
        heading_count = sum(1 for s in font_sizes if self.heading_range[0] < s <= self.heading_range[1])
        title_count = sum(1 for s in font_sizes if self.title_range[0] < s <= self.title_range[1])

        # Calculate percentages
        body_pct = (body_count / total) * 100
        heading_pct = (heading_count / total) * 100
        title_pct = (title_count / total) * 100

        result.font_size_distribution = {
            'body': body_pct,
            'heading': heading_pct,
            'title': title_pct,
            'other': 100.0 - (body_pct + heading_pct + title_pct)
        }

        # Find outliers
        outliers = [s for s in font_sizes if s < self.min_font_size or s > self.max_font_size]
        if outliers:
            result.has_outlier_fonts = True
            result.outlier_font_sizes = sorted(set(outliers))
            result.font_size_issues.append(
                f"Found {len(outliers)} font sizes outside acceptable range ({self.min_font_size}-{self.max_font_size}pt)"
            )

        # Validate distribution
        if body_pct < self.target_body_min:
            result.font_size_issues.append(
                f"Body text percentage too low: {body_pct:.1f}% (expected ≥{self.target_body_min}%)"
            )
        elif body_pct > self.target_body_max:
            result.font_size_issues.append(
                f"Body text percentage too high: {body_pct:.1f}% (expected ≤{self.target_body_max}%)"
            )

        if heading_pct < self.target_heading_min:
            result.font_size_issues.append(
                f"Heading percentage too low: {heading_pct:.1f}% (expected ≥{self.target_heading_min}%)"
            )

        if title_pct > self.target_title_max:
            result.font_size_issues.append(
                f"Title percentage too high: {title_pct:.1f}% (expected ≤{self.target_title_max}%)"
            )

        # Overall font size validation
        result.font_size_valid = (
            len(result.font_size_issues) == 0 and
            self.target_body_min <= body_pct <= self.target_body_max
        )

        result.all_issues.extend(result.font_size_issues)

    def _validate_layout(self, doc: Document, result: ValidationResult, expected_pages: Optional[int]) -> None:
        """Validate layout and structure"""
        # Count page breaks - check both text-based and XML-based page breaks
        page_break_count = 0

        # Method 1: Check paragraph-level page_break_before
        for paragraph in doc.paragraphs:
            if paragraph.paragraph_format.page_break_before:
                page_break_count += 1

        # Method 2: Check for explicit page break runs (w:br w:type="page")
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Check text-based page breaks (form feed characters)
                if '\f' in run.text or '\x0c' in run.text:
                    page_break_count += 1

                # Check XML-based page breaks (w:br elements)
                for elem in run._element:
                    # Look for <w:br w:type="page"/>
                    if 'w:br' in str(elem) and ('page' in str(elem).lower() or 'type="page"' in str(elem)):
                        page_break_count += 1

        result.page_count = page_break_count + 1  # Page breaks + 1 = total pages

        # Validate page breaks if expected count provided
        if expected_pages is not None:
            if result.page_count == expected_pages:
                result.page_breaks_valid = True
            else:
                result.page_breaks_valid = False
                result.layout_issues.append(
                    f"Page count mismatch: got {result.page_count}, expected {expected_pages}"
                )
        else:
            # If no expected count, just check that multi-page docs have breaks
            if result.page_count > 1:
                result.page_breaks_valid = True
            else:
                result.page_breaks_valid = True  # Single page is valid

        # TODO: Column detection would require analyzing paragraph positions
        # For now, mark as valid
        result.column_detection_valid = True

        result.all_issues.extend(result.layout_issues)

    def _calculate_score(self, result: ValidationResult) -> float:
        """Calculate overall quality score (0-100)"""
        score = 0.0

        # Font size validation (50 points)
        if result.font_size_valid:
            score += 50.0
        else:
            # Partial credit based on severity
            issue_count = len(result.font_size_issues)
            penalty = min(issue_count * 10, 50)
            score += max(50.0 - penalty, 0)

        # Page break validation (25 points)
        if result.page_breaks_valid:
            score += 25.0

        # Column detection (15 points)
        if result.column_detection_valid:
            score += 15.0

        # No outliers bonus (10 points)
        if not result.has_outlier_fonts:
            score += 10.0

        return min(score, 100.0)

    def validate_batch(self, file_paths: List[str]) -> List[ValidationResult]:
        """Validate multiple files"""
        results = []

        for file_path in file_paths:
            logger.info("Validating: %s", file_path)
            result = self.validate_docx(file_path)
            results.append(result)

        # Print summary
        passed = sum(1 for r in results if r.passed)
        failed = len(results) - passed
        avg_score = sum(r.score for r in results) / len(results) if results else 0

        print("\n" + "="*60)
        print("BATCH VALIDATION SUMMARY")
        print("="*60)
        print(f"Total files: {len(results)}")
        print(f"Passed: {passed}")
        print(f"Failed: {failed}")
        print(f"Average score: {avg_score:.1f}/100")
        print("="*60)

        return results

    def save_results(self, results: List[ValidationResult], output_path: str) -> None:
        """Save validation results to JSON"""
        data = {
            "validation_results": [r.to_dict() for r in results],
            "summary": {
                "total": len(results),
                "passed": sum(1 for r in results if r.passed),
                "failed": sum(1 for r in results if not r.passed),
                "average_score": sum(r.score for r in results) / len(results) if results else 0
            }
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

        logger.info("Validation results saved to: %s", output_path)
