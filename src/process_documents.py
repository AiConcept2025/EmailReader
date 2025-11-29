"""
Process original documents.
Process txt, pdf, rtf, images, foreign language documents documents
to word document
"""
import os
import logging
from docx import Document
from langdetect import detect  # type: ignore
from src.ocr import OCRProviderFactory
from src.pdf_image_ocr import ocr_pdf_image_to_doc
from src.document_analyzer import requires_ocr
from src.convert_to_docx import convert_pdf_to_docx
from src.file_utils import (
    delete_file,
    translate_document_to_english,
    convert_rtx_to_text
)
from src.file_utils import rename_file

# Get logger for this module
logger = logging.getLogger('EmailReader.DocProcessor')


class DocProcessor:
    """
    Class to process all type of docs and create
    collection of documents for processing by ChatGPT
    """

    def __init__(self, doc_path: str):
        """
        Initialise word document list
        and list for original documents
        """
        logger.debug("Initializing DocProcessor with path: %s", doc_path)
        # Word documents
        self.documents: list[str] = []
        # Original document
        self.original_documents: list[str] = []
        # document folder path
        self.docs_path: str = doc_path
        logger.debug("DocProcessor initialized successfully")

    # Process PDF documents
    def convert_pdf_payload_to_word(
            self,
            client: str,
            doc_name: str,
            payload: str) -> None:
        """
        Process PDF image or searchable doc and generate
        Word document
        Args:
        client:  client email
        doc_name: doc name '<email>+<doc name>'
        payload: content of the file
        """
        logger.info('Process PDF document %s', doc_name)
        file_name_no_ext, file_ext = os.path.splitext(doc_name)
        file_name = f'{client}+{file_name_no_ext}+original+original{file_ext}'
        file_path = os.path.join(self.docs_path, file_name)
        # Save original file in doc directory
        document = Document()
        document.add_paragraph(payload)
        document.save(file_path)

    # Process TIF document

    def convert_rtf_text_to_world(
            self,
            client: str,
            rtf_file_name: str,
            payload: str) -> None:
        """
        Converts a RTF text to a Word document.
        Args:
            client: client email
            txt_file_name: file name from attachment
            payload: payload from attachment
        """
        logger.info('Client %s convert TIF %s to Word document',
                    client, rtf_file_name)
        file_name_no_ext, _ = os.path.splitext(rtf_file_name)
        plain_text = convert_rtx_to_text(payload)

        rtf_file_path = os.path.join(
            self.docs_path,
            f'{client}+{file_name_no_ext}+original+original.tft')
        with open(rtf_file_path, 'w', encoding='utf-8') as fl:
            fl.write(payload)

        doc_file_path = os.path.join(
            self.docs_path,
            f'{client}+{file_name_no_ext}+original+original.doc')
        with open(doc_file_path, 'w', encoding='utf-8') as fl:
            fl.write(plain_text)

        # Check if foreign language
        if detect(plain_text) != 'en':
            translated_file_path = os.path.join(
                self.docs_path,
                f'{client}+{file_name_no_ext}+original+translated.doc')
            translate_document_to_english(doc_file_path, translated_file_path)

    def convert_rtf_file_to_world(
            self,
            client: str,
            rtf_file_name: str
    ) -> None:
        """
            Converts a RTF file to a Word document.
            Args:
                client: client email
                txt_file_name: file name from google drive
        """
        logger.info('Client %s convert TIF %s to Word document',
                    client, rtf_file_name)
        with open(rtf_file_name, 'r', encoding='utf-8') as fl:
            rtf_text = fl.read()

        file_name_no_ext, _ = os.path.splitext(rtf_file_name)
        plain_text = convert_rtx_to_text(rtf_text)

        rtf_file_path = os.path.join(
            self.docs_path,
            f'{client}+{file_name_no_ext}+original+original.tft')
        with open(rtf_file_path, 'w', encoding='utf-8') as fl:
            fl.write(rtf_text)

        doc_file_path = os.path.join(
            self.docs_path,
            f'{client}+{file_name_no_ext}+original+original.doc')
        with open(doc_file_path, 'w', encoding='utf-8') as fl:
            fl.write(plain_text)

        # Check if foreign language
        if detect(plain_text) != 'en':
            translated_file_path = os.path.join(
                self.docs_path,
                f'{client}+{file_name_no_ext}+original+translated.doc')
            translate_document_to_english(doc_file_path, translated_file_path)

    # Process plain text
    def convert_plain_text_to_word(
            self,
            client: str,
            txt_file_name: str,
            payload: str) -> None:
        """
        Converts a plain text to a Word document.
        Args:
            client: client email
            txt_file_name: file name from attachment
            payload: payload from attachment
        """
        logger.info('Convert txt to word %s', txt_file_name)
        file_name_no_ext, _ = os.path.splitext(txt_file_name)
        file_path = os.path.join(
            self.docs_path,
            f'{client}+{file_name_no_ext}+original+original.txt')
        doc_file_path = os.path.join(
            self.docs_path,
            f'{client}+{file_name_no_ext}+original+original.doc')
        with open(file_path, 'w', encoding='utf-8') as fl:
            fl.write(payload)
        document = Document()
        document.add_paragraph(payload)
        document.save(doc_file_path)

        # Check if foreign language
        if detect(payload) != 'en':
            translated_file_path = os.path.join(
                self.docs_path,
                f'{client}+{file_name_no_ext}+original+translated.doc')
            translate_document_to_english(doc_file_path, translated_file_path)

    def convert_plain_text_file_to_word(
            self,
            client: str,
            txt_file_name: str
    ) -> None:
        """
        Converts a plain text file to a Word document.
        Args:
            client: client email
            txt_file_name: file name from attachment
            payload: payload from attachment
        """
        logger.info('Convert txt file to word %s', txt_file_name)
        file_name_no_ext, _ = os.path.splitext(txt_file_name)
        file_path = os.path.join(
            self.docs_path,
            f'{client}+{file_name_no_ext}+original+original.txt')
        with open(file_path, 'r', encoding='utf-8') as fl:
            plain_text = fl.read()
        doc_file_path = os.path.join(
            self.docs_path,
            f'{client}+{file_name_no_ext}+original+original.doc')
        document = Document()
        document.add_paragraph(plain_text)
        document.save(doc_file_path)

        # Check if foreign language
        if detect(plain_text) != 'en':
            translated_file_path = os.path.join(
                self.docs_path,
                f'{client}+{file_name_no_ext}+original+translated.doc')
            translate_document_to_english(doc_file_path, translated_file_path)

    # Process Word document
    def process_word_load(
            self,
            client: str,
            word_file_name: str,
            payload: str
    ) -> None:
        """
        Process Word document text to Word file
        Args:
            client: client email
            txt_file_name: file name from attachment
            payload: payload from attachment
        """
        logger.info('Process Word doc text to file %s', word_file_name)
        if detect(payload) == 'en':
            word_file_name = f'{client}+{word_file_name}+original+english.docx'
            word_file_path = os.path.join(self.docs_path, word_file_name)
            with open(word_file_path, 'w', encoding='utf-8') as fl:
                fl.write(payload)
        else:
            word_file_name = f'{client}+{word_file_name}+original+foreign.docx'
            word_file_path = os.path.join(self.docs_path, word_file_name)
            with open(word_file_path, 'w', encoding='utf-8') as fl:
                fl.write(payload)
            translated_file_name = (
                f'{client}+{word_file_name}+original+translated.docx'
            )
            translated_file_path = os.path.join(
                self.docs_path, translated_file_name)
            translate_document_to_english(word_file_path, translated_file_path)

    def process_word_file(
            self,
            client: str,  # Keep parameter for backward compatibility but don't use it
            file_name: str,
            document_folder: str,
            target_lang: str | None = None
    ) -> tuple[str, str, str, str]:
        """
        Process Word document text to Word file and translate it if needed.
        Args:
            client: client email (no longer used in file naming)
            file_name: new file in google drive inbox sub folder
            document_folder: folder to temp save processed documents
        """
        logger.info('Starting process_word_file() for: %s', file_name)
        logger.debug("Parameters - client: %s, folder: %s, target_lang: %s",
                     client, document_folder, target_lang)

        try:
            file_path = os.path.join(document_folder, file_name)
            logger.debug("Full file path: %s", file_path)

            if not os.path.exists(file_path):
                logger.error("File not found: %s", file_path)
                raise FileNotFoundError(f"File not found: {file_path}")

            file_name_no_ext, file_ext = os.path.splitext(file_name)
            logger.debug("File name without ext: %s, extension: %s",
                         file_name_no_ext, file_ext)

            logger.debug("Loading Word document")
            document = Document(file_path)

            # Check if document language is English
            logger.debug("Extracting text from document")
            full_text: list[str] = []
            for paragraph in document.paragraphs:
                full_text.append(paragraph.text)
            text = '\n'.join(full_text)
            text_length = len(text)
            logger.debug("Extracted text length: %d characters", text_length)

            if text_length < 10:
                logger.warning(
                    "Document has very little text (%d chars)", text_length)

            logger.debug("Detecting language")
            detected_lang = detect(text)
            logger.info("Detected language: %s", detected_lang)

            # If English, rename file with +english
            if detected_lang == 'en':
                logger.info("Document is in English, no translation needed")
                # Do not prefix with client here; caller will prepend email
                new_file_name = f'{file_name_no_ext}+english{file_ext}'
                new_file_path = os.path.join(document_folder, new_file_name)
                logger.debug("Renaming to: %s", new_file_name)
                rename_file(file_path, new_file_path)
                original_file_name = new_file_name
                original_file_path = new_file_path
            else:
                # If not English, translate it and rename file with +translated
                logger.info(
                    "Document is in %s, translation required", detected_lang)
                # Rename original file with +original
                # Do not prefix with client here; caller will prepend email
                original_file_name = f'{file_name_no_ext}+original{file_ext}'
                original_file_path = os.path.join(
                    document_folder, original_file_name)
                logger.debug("Renaming original to: %s", original_file_name)
                rename_file(file_path, original_file_path)

                # Translate document (optionally to target_lang)
                new_file_name = f'{file_name_no_ext}+translated{file_ext}'
                new_file_path = os.path.join(document_folder, new_file_name)
                logger.info("Translating document from %s to %s",
                            detected_lang, target_lang or 'en')
                translate_document_to_english(
                    original_file_path, new_file_path, target_lang)
                logger.info("Translation completed: %s", new_file_name)

            logger.info("process_word_file() completed successfully")
            logger.debug("Returning: new_file=%s, original_file=%s",
                         new_file_name, original_file_name)

            return (
                new_file_path,
                new_file_name,
                original_file_name,
                original_file_path)

        except Exception as e:
            logger.error("Error in process_word_file() for %s: %s",
                         file_name, e, exc_info=True)
            raise

    def _process_with_ocr_provider(
            self,
            input_file: str,
            output_file: str
    ) -> None:
        """
        Process document with configured OCR provider, with automatic fallback.

        Args:
            input_file: Path to input file (PDF or image)
            output_file: Path to save output DOCX file
        """
        from src.config import load_config
        from src.ocr.default_provider import DefaultOCRProvider

        try:
            # Load config and get provider
            config = load_config()
            provider_name = config.get('ocr', {}).get('provider', 'default')

            logger.info(f"Using OCR provider: {provider_name}")
            ocr_provider = OCRProviderFactory.get_provider(config)
            ocr_provider.process_document(input_file, output_file)
            logger.info(
                f"OCR completed successfully with {provider_name} provider")

        except Exception as e:
            logger.warning(
                (f"Primary OCR provider failed: {e}. "
                 "Falling back to default Tesseract OCR")
            )

            # Fallback to default provider
            try:
                fallback_provider = DefaultOCRProvider({})
                fallback_provider.process_document(input_file, output_file)
                logger.info("Fallback OCR completed successfully")
            except Exception as fallback_error:
                logger.error(f"Fallback OCR also failed: {fallback_error}")
                raise RuntimeError(
                    f"Both primary and fallback OCR failed. "
                    f"Primary: {e}, Fallback: {fallback_error}"
                ) from fallback_error

    def convert_pdf_file_to_word(
            self,
            client: str,  # Keep parameter for backward compatibility but don't use it
            file_name: str,
            document_folder: str,
            metadata: dict | None = None,
            target_lang: str | None = None
    ) -> tuple[str, str, str, str]:
        """
        Convert PDF file (image or text) to word document
        Args:
            client: client email (no longer used in file naming)
            file_name: new file in google drive inbox sub folder
            document_folder: folder to temp save processed documents
        """
        logger.info('Starting convert_pdf_file_to_word() for: %s', file_name)
        logger.debug(
            "Parameters - client: %s, folder: %s, target_lang: %s",
            client,
            document_folder,
            target_lang)

        try:
            file_path = os.path.join(document_folder, file_name)
            logger.debug("Full PDF path: %s", file_path)

            if not os.path.exists(file_path):
                logger.error("PDF file not found: %s", file_path)
                raise FileNotFoundError(f"File not found: {file_path}")

            file_name_no_ext, file_ext = os.path.splitext(file_name)
            logger.debug("File name without ext: %s, extension: %s",
                         file_name_no_ext, file_ext)

            # CHANGED: Removed {client}+ prefix
            original_file_name = f'{file_name_no_ext}+original{file_ext}'
            original_file_path = os.path.join(
                document_folder, original_file_name)

            # Rename file to original
            logger.debug("Renaming PDF to original: %s", original_file_name)
            rename_file(file_path, original_file_path)

            docx_file_path = os.path.join(
                document_folder, f'{file_name_no_ext}.docx')
            logger.debug("Target DOCX path: %s", docx_file_path)

            # Determine if OCR is needed
            logger.info("Analyzing PDF to determine if OCR is required")
            needs_ocr = requires_ocr(original_file_path)
            logger.info(
                "PDF analysis complete: %s",
                "OCR required" if needs_ocr else "Searchable text found")

            if needs_ocr:
                logger.info("Processing document with OCR provider")
                ocr_method = metadata.get('ocr_method', 'default')
                if ocr_method == 'default':
                    ocr_pdf_image_to_doc(
                        original_file_path,
                        docx_file_path)
                else:
                    self._process_with_ocr_provider(
                        original_file_path,
                        docx_file_path)

                logger.info("OCR processing completed")
            else:
                logger.info(
                    "Converting searchable PDF to DOCX (no OCR needed)")
                convert_pdf_to_docx(original_file_path, docx_file_path)
                logger.info("PDF to DOCX conversion completed")

            # CHANGED: Removed {client}+ prefix
            new_file_name = f'{file_name_no_ext}+translated.docx'
            new_file_path = os.path.join(document_folder, new_file_name)
            logger.info("Translating PDF content to %s", target_lang or 'en')
            translate_document_to_english(
                docx_file_path, new_file_path, target_lang)
            logger.info("Translation completed: %s", new_file_name)

            logger.debug("Cleaning up temporary DOCX file: %s", docx_file_path)
            delete_file(docx_file_path)

            logger.info("convert_pdf_file_to_word() completed successfully")
            logger.debug("Returning: new_file=%s, original_file=%s",
                         new_file_name, original_file_name)

            return (
                new_file_path,
                new_file_name,
                original_file_name,
                original_file_path)

        except Exception as e:
            logger.error(
                "Error in convert_pdf_file_to_word() for %s: %s", file_name, e, exc_info=True)
            raise
