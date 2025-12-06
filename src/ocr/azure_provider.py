"""
Azure OCR Provider

Implements Azure Document Intelligence Read API integration for OCR.
"""

import os
import time
import logging
from typing import Dict, Any, List, Tuple
from pathlib import Path

import pdfplumber
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from azure.core.exceptions import HttpResponseError

from src.ocr.base_provider import BaseOCRProvider
from src.models.paragraph import Paragraph, TextSpan

logger = logging.getLogger('EmailReader.OCR.Azure')


class AzureOCRProvider(BaseOCRProvider):
    """
    Azure Document Intelligence OCR provider.

    Uses Azure's Read API for intelligent document analysis with
    layout preservation and high accuracy OCR.
    """

    def __init__(self, config: Dict[str, Any]):
        """
        Initialize the Azure OCR provider.

        Args:
            config: Configuration dictionary with 'endpoint' and 'api_key'

        Raises:
            ValueError: If required config values are missing
        """
        super().__init__(config)

        self.endpoint = config.get('endpoint')
        self.api_key = config.get('api_key')

        if not self.endpoint or not self.api_key:
            raise ValueError(
                "Azure OCR provider requires 'endpoint' and 'api_key' in configuration"
            )

        logger.info("Initializing Azure Document Intelligence client")
        logger.debug("Azure endpoint: %s", self.endpoint)

        # Initialize paragraph extraction flag (will be set by OCRProviderFactory if needed)
        # Default: True to use paragraph-based extraction by default
        # This preserves document structure and formatting
        self.use_paragraph_extraction = True

        try:
            self.client = DocumentAnalysisClient(
                endpoint=self.endpoint,
                credential=AzureKeyCredential(self.api_key)
            )
            logger.info("Azure OCR provider initialized successfully")
        except Exception as e:
            logger.error("Failed to initialize Azure client: %s", e)
            raise

    def process_document(self, input_path: str, output_path: str, use_paragraph_extraction: bool = None) -> None:
        """
        Process a document with Azure OCR and save the result.

        Detects searchable vs scanned pages and processes accordingly.

        Args:
            input_path: Path to input PDF file
            output_path: Path to save output DOCX file
            use_paragraph_extraction: If True, extract paragraphs with formatting metadata.
                                     If False, use line-based extraction.
                                     If None (default), uses the instance's use_paragraph_extraction attribute.

        Raises:
            FileNotFoundError: If input file doesn't exist
            RuntimeError: If OCR processing fails
        """
        # Use instance attribute if no explicit parameter provided
        if use_paragraph_extraction is None:
            use_paragraph_extraction = self.use_paragraph_extraction

        logger.info("Processing document with Azure OCR: %s (paragraph_extraction=%s)",
                   os.path.basename(input_path), use_paragraph_extraction)
        logger.debug("Input: %s", input_path)
        logger.debug("Output: %s", output_path)

        if not os.path.exists(input_path):
            logger.error("Input file not found: %s", input_path)
            raise FileNotFoundError(f"File not found: {input_path}")

        try:
            # Detect which pages need OCR
            logger.info("Analyzing PDF pages for searchability")
            page_searchability = self._detect_page_searchability(input_path)
            total_pages = len(page_searchability)
            searchable_count = sum(1 for is_searchable in page_searchability if is_searchable)
            ocr_count = total_pages - searchable_count

            logger.info("PDF analysis: %d total pages, %d searchable, %d need OCR",
                       total_pages, searchable_count, ocr_count)

            # Extract all pages content
            pages_content: List[str] = []

            with pdfplumber.open(input_path) as pdf:
                for page_num, (page, is_searchable) in enumerate(zip(pdf.pages, page_searchability), 1):
                    logger.debug("Processing page %d/%d (searchable: %s)",
                               page_num, total_pages, is_searchable)

                    if is_searchable:
                        # Extract text directly
                        text = page.extract_text() or ""
                        logger.debug("Page %d: extracted %d characters directly",
                                   page_num, len(text))
                        pages_content.append(text)
                    else:
                        # This page needs OCR - we'll OCR the entire document
                        # Break out and use Azure for all pages
                        logger.info("Found non-searchable page %d - using Azure OCR for entire document",
                                  page_num)
                        break
                else:
                    # All pages are searchable
                    if use_paragraph_extraction:
                        # User requested paragraph extraction - use Azure OCR even for searchable PDFs
                        # to get proper paragraph structure and formatting
                        logger.info("All pages searchable, but using Azure OCR for paragraph extraction")
                        # Continue to Azure OCR processing below (don't return)
                    else:
                        # No paragraph extraction requested - save pages as-is
                        logger.info("All pages are searchable - saving without OCR")
                        self._save_as_docx(pages_content, output_path)
                        return

            # At least one page needs OCR - use Azure for entire document
            logger.info("Reading PDF file for Azure OCR")
            with open(input_path, 'rb') as f:
                pdf_bytes = f.read()

            file_size_kb = len(pdf_bytes) / 1024
            logger.debug("PDF file size: %.2f KB", file_size_kb)

            # Perform OCR with Azure - choose extraction mode
            if use_paragraph_extraction:
                logger.info("Using paragraph extraction mode")
                paragraphs = self._extract_paragraphs_with_formatting(pdf_bytes)
                self._save_paragraphs_to_docx(paragraphs, output_path)
            else:
                logger.info("Using line-based extraction mode (backward compatible)")
                ocr_result = self._ocr_with_azure(pdf_bytes)
                self._save_as_docx(ocr_result, output_path)

            if os.path.exists(output_path):
                output_size = os.path.getsize(output_path) / 1024
                logger.info("Azure OCR completed: %s (%.2f KB)",
                           os.path.basename(output_path), output_size)
            else:
                raise RuntimeError("OCR processing failed to create output file")

        except FileNotFoundError:
            raise
        except Exception as e:
            logger.error("Error during Azure OCR processing: %s", e, exc_info=True)
            raise RuntimeError(f"Azure OCR processing failed: {e}")

    def is_pdf_searchable(self, pdf_path: str) -> bool:
        """
        Check if a PDF contains searchable text.

        Uses pdfplumber to check if >50% of pages have >50 characters.

        Args:
            pdf_path: Path to PDF file

        Returns:
            True if PDF has extractable text, False otherwise
        """
        logger.debug("Checking if PDF is searchable: %s", os.path.basename(pdf_path))

        try:
            with pdfplumber.open(pdf_path) as pdf:
                pages_with_text = 0
                total_pages = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text and len(text.strip()) > 50:
                        pages_with_text += 1
                        logger.debug("Page %d has sufficient text", page_num)

                is_searchable = pages_with_text > (total_pages * 0.5)
                logger.debug("PDF searchable: %s (%d/%d pages with text)",
                           is_searchable, pages_with_text, total_pages)
                return is_searchable

        except Exception as e:
            logger.error("Error checking PDF searchability: %s", e)
            raise

    def _detect_page_searchability(self, pdf_path: str) -> List[bool]:
        """
        Analyze each page individually to determine if it needs OCR.

        Args:
            pdf_path: Path to PDF file

        Returns:
            List of booleans indicating if each page is searchable
        """
        logger.debug("Detecting searchability for each page")

        try:
            page_searchability: List[bool] = []

            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    char_count = len(text.strip()) if text else 0
                    is_searchable = char_count > 50

                    logger.debug("Page %d: %d characters, searchable: %s",
                               page_num, char_count, is_searchable)
                    page_searchability.append(is_searchable)

            return page_searchability

        except Exception as e:
            logger.error("Error detecting page searchability: %s", e)
            raise

    def _ocr_with_azure(self, pdf_bytes: bytes, max_retries: int = 3) -> List[str]:
        """
        Call Azure Read API to perform OCR.

        Args:
            pdf_bytes: PDF file content as bytes
            max_retries: Maximum number of retry attempts

        Returns:
            List of text content for each page

        Raises:
            RuntimeError: If OCR fails after retries
        """
        logger.info("Starting Azure Document Intelligence Read API call")
        logger.debug("Document size: %d bytes", len(pdf_bytes))

        for attempt in range(1, max_retries + 1):
            try:
                logger.debug("Azure API attempt %d/%d", attempt, max_retries)

                start_time = time.time()

                # Begin analyzing document
                logger.debug("Calling begin_analyze_document with prebuilt-read model")
                poller = self.client.begin_analyze_document(
                    "prebuilt-read",
                    document=pdf_bytes
                )

                logger.info("Azure analysis started, waiting for completion...")

                # Wait for completion
                result = poller.result()

                duration = time.time() - start_time
                logger.info("Azure analysis completed in %.2f seconds", duration)
                logger.debug("Result contains %d pages", len(result.pages))

                # Extract text from each page
                pages_content: List[str] = []

                for page_num, page in enumerate(result.pages, 1):
                    logger.debug("Extracting text from page %d/%d",
                               page_num, len(result.pages))

                    # Get all lines on this page
                    page_lines = [
                        line.content
                        for line in result.pages[page_num - 1].lines
                    ]

                    page_text = "\n".join(page_lines)
                    char_count = len(page_text)

                    logger.debug("Page %d: extracted %d characters",
                               page_num, char_count)

                    pages_content.append(page_text)

                total_chars = sum(len(p) for p in pages_content)
                logger.info("Azure OCR extracted %d characters from %d pages",
                           total_chars, len(pages_content))

                return pages_content

            except HttpResponseError as e:
                logger.error("Azure API HTTP error (attempt %d/%d): %s",
                           attempt, max_retries, e)

                if attempt < max_retries:
                    wait_time = 2 ** attempt  # Exponential backoff
                    logger.info("Retrying in %d seconds...", wait_time)
                    time.sleep(wait_time)
                else:
                    logger.error("Max retries reached, giving up")
                    raise RuntimeError(f"Azure OCR failed after {max_retries} attempts: {e}")

            except Exception as e:
                logger.error("Unexpected error during Azure OCR: %s", e, exc_info=True)
                raise RuntimeError(f"Azure OCR failed: {e}")

        raise RuntimeError("Azure OCR failed: max retries exceeded")

    def _merge_split_paragraphs(self, paragraphs: List[Paragraph]) -> List[Paragraph]:
        """
        Merge paragraphs that were split mid-sentence by PDF layout.

        Azure OCR detects visual lines as separate paragraphs. This method merges
        consecutive paragraphs that are clearly part of the same sentence (e.g.,
        lines that don't end with proper punctuation).

        Args:
            paragraphs: List of paragraphs extracted from Azure OCR

        Returns:
            List of paragraphs with split lines merged

        Logic:
            - If a paragraph doesn't end with sentence-ending punctuation (. ! ? : ;)
            - And the next paragraph is on the same page
            - And they have the same role
            - Then merge them into a single paragraph
        """
        if not paragraphs or len(paragraphs) < 2:
            return paragraphs

        logger.info("Merging split paragraphs (before: %d paragraphs)", len(paragraphs))

        merged = []
        i = 0

        while i < len(paragraphs):
            current = paragraphs[i]

            # Collect consecutive paragraphs to merge
            paragraphs_to_merge = [current]

            # Check if we should merge with next paragraphs
            while i + 1 < len(paragraphs):
                next_para = paragraphs[i + 1]

                # Conditions for merging:
                # 1. Current paragraph doesn't end with sentence-ending punctuation
                # 2. Both on same page
                # 3. Same role (don't merge titles with paragraphs, etc.)
                # 4. Next paragraph is not a list item (list items are separate)

                current_text = current.content.rstrip()

                # Define sentence-ending punctuation
                sentence_endings = ('.', '!', '?', ':', ';', '。', '！', '？')  # Include some unicode punctuation

                should_merge = (
                    current_text and
                    not current_text.endswith(sentence_endings) and
                    current.page == next_para.page and
                    current.role == next_para.role and
                    not next_para.is_list_item
                )

                if should_merge:
                    logger.debug("Merging paragraph %d with %d (no sentence ending: '%s')",
                               i, i + 1, current_text[-20:] if len(current_text) > 20 else current_text)
                    paragraphs_to_merge.append(next_para)
                    current = next_para  # Move to next for continued checking
                    i += 1
                else:
                    break

            # Create merged paragraph
            if len(paragraphs_to_merge) > 1:
                # Merge content with spaces
                merged_content = ' '.join(p.content.strip() for p in paragraphs_to_merge)

                # Combine spans from all paragraphs
                merged_spans = []
                for p in paragraphs_to_merge:
                    merged_spans.extend(p.spans)

                # Create merged paragraph using first paragraph's metadata
                first = paragraphs_to_merge[0]
                merged_para = Paragraph(
                    content=merged_content,
                    page=first.page,
                    role=first.role,
                    spans=merged_spans,
                    bounding_box=first.bounding_box,  # Keep first paragraph's bbox
                    is_list_item=first.is_list_item,
                    list_marker=first.list_marker
                )

                logger.debug("Merged %d paragraphs into one (%d chars)",
                           len(paragraphs_to_merge), len(merged_content))
                merged.append(merged_para)
            else:
                # No merge needed
                merged.append(current)

            i += 1

        logger.info("Merging complete (after: %d paragraphs, reduced by %d)",
                   len(merged), len(paragraphs) - len(merged))

        return merged

    def _extract_paragraphs_with_formatting(self, pdf_bytes: bytes, max_retries: int = 3) -> List[Paragraph]:
        """
        Extract paragraphs with formatting from PDF using Azure Document Intelligence.

        Uses Azure's paragraph extraction instead of line-based extraction to preserve
        document structure and formatting information.

        Args:
            pdf_bytes: PDF file content as bytes
            max_retries: Maximum number of retry attempts

        Returns:
            List of Paragraph objects with content, role, and formatting metadata

        Raises:
            RuntimeError: If OCR fails after retries
        """
        logger.info("Starting Azure Document Intelligence paragraph extraction")
        logger.debug("Document size: %d bytes", len(pdf_bytes))

        for attempt in range(1, max_retries + 1):
            try:
                logger.debug("Azure API attempt %d/%d", attempt, max_retries)

                start_time = time.time()

                # Begin analyzing document
                logger.debug("Calling begin_analyze_document with prebuilt-read model")
                poller = self.client.begin_analyze_document(
                    "prebuilt-read",
                    document=pdf_bytes
                )

                logger.info("Azure analysis started, waiting for completion...")

                # Wait for completion
                result = poller.result()

                duration = time.time() - start_time
                logger.info("Azure analysis completed in %.2f seconds", duration)

                # Check if paragraphs are available
                if not hasattr(result, 'paragraphs') or not result.paragraphs:
                    logger.warning("No paragraphs found in Azure response, falling back to empty list")
                    return []

                logger.debug("Result contains %d paragraphs", len(result.paragraphs))

                # Extract paragraphs with formatting
                paragraphs: List[Paragraph] = []

                for para_num, azure_para in enumerate(result.paragraphs, 1):
                    logger.debug("Processing paragraph %d/%d", para_num, len(result.paragraphs))

                    # Extract content (required)
                    raw_content = azure_para.content if hasattr(azure_para, 'content') else ""

                    # BUG-001 FIX: Clean embedded newlines that break sentences mid-phrase
                    # Azure returns paragraph content with embedded \n characters from PDF layout
                    # Strategy: Replace single newlines with spaces, preserve double newlines (rare)
                    newline_count = raw_content.count('\n')
                    if newline_count > 0:
                        logger.warning("[FIX BUG-001] Paragraph %d contains %d embedded newlines - cleaning",
                                      para_num, newline_count)
                        logger.debug("[FIX BUG-001] Raw content preview: %s",
                                    raw_content[:100].replace('\n', '\\n'))

                        # Preserve double newlines (paragraph separators), replace single newlines
                        content = raw_content.replace('\n\n', '§§PARA§§')  # Temporary marker
                        content = content.replace('\n', ' ')  # Replace line breaks with spaces
                        content = content.replace('§§PARA§§', '\n\n')  # Restore paragraph breaks

                        # Clean up excessive whitespace
                        content = ' '.join(content.split())

                        logger.debug("[FIX BUG-001] Cleaned content preview: %s", content[:100])
                    else:
                        content = raw_content

                    # Extract page number (default to 1 if not available)
                    page = 1
                    if hasattr(azure_para, 'page_number'):
                        page = azure_para.page_number
                    elif hasattr(azure_para, 'bounding_regions') and azure_para.bounding_regions:
                        page = azure_para.bounding_regions[0].page_number

                    # Extract role (default to 'paragraph' if not available)
                    if hasattr(azure_para, 'role') and azure_para.role:
                        role = azure_para.role
                    else:
                        role = 'paragraph'

                    logger.debug("Paragraph %d: page=%d, role=%s, length=%d",
                               para_num, page, role, len(content))

                    # Extract bounding box if available
                    bounding_box = None
                    if hasattr(azure_para, 'bounding_regions') and azure_para.bounding_regions:
                        try:
                            region = azure_para.bounding_regions[0]
                            if hasattr(region, 'polygon') and region.polygon:
                                # Extract coordinates from polygon points
                                polygon = region.polygon
                                if len(polygon) >= 4:
                                    # Calculate bounding box from polygon
                                    x_coords = [point.x for point in polygon]
                                    y_coords = [point.y for point in polygon]
                                    x_min = min(x_coords)
                                    y_min = min(y_coords)
                                    x_max = max(x_coords)
                                    y_max = max(y_coords)

                                    bounding_box = {
                                        'x': x_min,
                                        'y': y_min,
                                        'width': x_max - x_min,
                                        'height': y_max - y_min
                                    }
                                    logger.debug("Extracted bounding box: %s", bounding_box)
                        except Exception as e:
                            logger.debug("Could not extract bounding box: %s", e)

                    # Extract text spans with formatting (if available)
                    spans: List[TextSpan] = []
                    if hasattr(azure_para, 'spans') and azure_para.spans:
                        for span in azure_para.spans:
                            # Azure Read API may not provide detailed formatting
                            # For now, create basic spans
                            span_text = content  # Simplified - full content
                            text_span = TextSpan(
                                text=span_text,
                                is_bold=False,  # Not available in Read API
                                is_italic=False,  # Not available in Read API
                                font_size=None  # Not available in Read API
                            )
                            spans.append(text_span)
                            logger.debug("Created text span with %d characters", len(span_text))
                    else:
                        # No spans - create default span with full content
                        spans = [TextSpan(text=content)]

                    # Detect list items
                    is_list_item = role == 'listItem'
                    list_marker = None
                    if is_list_item and content:
                        # Extract common list markers
                        for marker in ['•', '◦', '▪', '-', '*']:
                            if content.strip().startswith(marker):
                                list_marker = marker
                                break
                        # Check for numbered lists
                        if not list_marker and content[0].isdigit():
                            idx = content.find('.')
                            if idx > 0 and idx < 5:
                                list_marker = content[:idx+1]

                    # Create paragraph object
                    paragraph = Paragraph(
                        content=content,
                        page=page,
                        role=role,
                        spans=spans,
                        bounding_box=bounding_box,
                        is_list_item=is_list_item,
                        list_marker=list_marker
                    )

                    paragraphs.append(paragraph)

                # Merge split paragraphs (lines broken mid-sentence by PDF layout)
                paragraphs = self._merge_split_paragraphs(paragraphs)

                # BUG-005 FIX: Add diagnostic logging summary
                role_counts = {}
                total_embedded_newlines = 0
                for para in paragraphs:
                    role_counts[para.role] = role_counts.get(para.role, 0) + 1
                    # Count any remaining embedded newlines (should be 0 after cleaning)
                    total_embedded_newlines += para.content.count('\n')

                total_chars = sum(len(p.content) for p in paragraphs)
                logger.info("Extracted %d paragraphs, %d total characters", len(paragraphs), total_chars)
                logger.info("Paragraph roles detected: %s", role_counts)

                # BUG-005 FIX: Report embedded newlines status
                if total_embedded_newlines > 0:
                    logger.error("[FIX BUG-005] WARNING: %d embedded newlines remain after cleaning!",
                                total_embedded_newlines)
                else:
                    logger.info("[FIX BUG-005] All embedded newlines cleaned successfully")

                return paragraphs

            except HttpResponseError as e:
                logger.error("Azure API HTTP error (attempt %d/%d): %s",
                           attempt, max_retries, e)

                if attempt < max_retries:
                    wait_time = 2 ** attempt  # Exponential backoff
                    logger.info("Retrying in %d seconds...", wait_time)
                    time.sleep(wait_time)
                else:
                    logger.error("Max retries reached, giving up")
                    raise RuntimeError(f"Azure paragraph extraction failed after {max_retries} attempts: {e}")

            except Exception as e:
                logger.error("Unexpected error during Azure paragraph extraction: %s", e, exc_info=True)
                raise RuntimeError(f"Azure paragraph extraction failed: {e}")

        raise RuntimeError("Azure paragraph extraction failed: max retries exceeded")

    def _save_as_docx(self, pages_content: List[str], output_path: str) -> None:
        """
        Save extracted text as a DOCX file with page breaks.

        Args:
            pages_content: List of text content for each page
            output_path: Path to save DOCX file
        """
        logger.debug("Saving %d pages to DOCX: %s",
                    len(pages_content), os.path.basename(output_path))

        try:
            document = Document()

            for page_num, page_text in enumerate(pages_content, 1):
                logger.debug("Adding page %d to document (%d characters)",
                           page_num, len(page_text))

                # Add page content
                if page_text.strip():
                    # CRITICAL FIX: Clean embedded newlines from extracted text
                    # pdfplumber.extract_text() returns text with embedded \n characters
                    # that break sentences mid-phrase when read by python-docx
                    clean_text = ' '.join(page_text.split('\n'))
                    clean_text = ' '.join(clean_text.split())  # Also clean excessive whitespace

                    paragraph = document.add_paragraph(clean_text)
                    # Set font size
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                else:
                    logger.debug("Page %d is empty", page_num)
                    document.add_paragraph("")

                # Add page break (except after last page)
                if page_num < len(pages_content):
                    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                    logger.debug("Added page break after page %d", page_num)

            logger.debug("Saving DOCX file to: %s", output_path)
            document.save(output_path)

            logger.info("DOCX file saved successfully: %s", os.path.basename(output_path))

        except Exception as e:
            logger.error("Error saving DOCX: %s", e, exc_info=True)
            raise

    def _save_paragraphs_to_docx(self, paragraphs: List[Paragraph], output_path: str) -> None:
        """
        Save paragraphs with formatting as a DOCX file.

        Applies appropriate styles based on paragraph roles and preserves formatting
        information from text spans. Adds page breaks between pages.

        Args:
            paragraphs: List of Paragraph objects with content and formatting
            output_path: Path to save DOCX file

        Raises:
            Exception: If saving DOCX fails
        """
        logger.info("Saving %d paragraphs to DOCX with formatting: %s",
                   len(paragraphs), os.path.basename(output_path))

        try:
            document = Document()
            current_page = 1

            for para_num, para in enumerate(paragraphs, 1):
                logger.debug("Adding paragraph %d/%d (page=%d, role=%s)",
                           para_num, len(paragraphs), para.page, para.role)

                # Add page break if we've moved to a new page
                if para.page > current_page:
                    logger.debug("Adding page break before page %d", para.page)
                    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                    current_page = para.page

                # Determine style based on role
                if para.role == 'title':
                    # Title - use Heading level 0 (Title style)
                    # SAFETY: Clean embedded newlines before adding
                    clean_content = ' '.join(para.content.split('\n'))
                    clean_content = ' '.join(clean_content.split())
                    doc_para = document.add_heading(clean_content, level=0)
                    logger.debug("Added title paragraph")
                    # Add spacing after title
                    if para_num < len(paragraphs):
                        document.add_paragraph("")
                        logger.debug("Added spacing after title")

                elif para.role == 'heading':
                    # Heading - use Heading level 1
                    # SAFETY: Clean embedded newlines before adding
                    clean_content = ' '.join(para.content.split('\n'))
                    clean_content = ' '.join(clean_content.split())
                    doc_para = document.add_heading(clean_content, level=1)
                    logger.debug("Added heading paragraph")
                    # Add spacing after heading
                    if para_num < len(paragraphs):
                        document.add_paragraph("")
                        logger.debug("Added spacing after heading")

                elif para.role == 'listItem' or para.is_list_item:
                    # List item - add with marker if available
                    # SAFETY: Clean embedded newlines before adding
                    clean_para_content = ' '.join(para.content.split('\n'))
                    clean_para_content = ' '.join(clean_para_content.split())

                    if para.list_marker:
                        content_with_marker = f"{para.list_marker} {clean_para_content.lstrip(para.list_marker).strip()}"
                    else:
                        content_with_marker = clean_para_content

                    doc_para = document.add_paragraph(content_with_marker)
                    logger.debug("Added list item paragraph with marker: %s", para.list_marker)
                    # Add spacing after list item
                    if para_num < len(paragraphs):
                        document.add_paragraph("")
                        logger.debug("Added spacing after list item")

                else:
                    # Regular paragraph
                    doc_para = document.add_paragraph()

                    # Add content with formatting from spans if available
                    if para.spans and len(para.spans) > 0:
                        for span in para.spans:
                            # SAFETY: Final cleanup of any embedded newlines before adding to DOCX
                            # This ensures no newlines slip through regardless of code path
                            clean_text = ' '.join(span.text.split('\n'))
                            clean_text = ' '.join(clean_text.split())  # Clean excessive whitespace
                            run = doc_para.add_run(clean_text)

                            # Apply formatting
                            if span.is_bold:
                                run.bold = True
                            if span.is_italic:
                                run.italic = True
                            if span.font_size:
                                run.font.size = Pt(span.font_size)

                            logger.debug("Added formatted span: bold=%s, italic=%s, size=%s",
                                       span.is_bold, span.is_italic, span.font_size)
                    else:
                        # No spans - add plain text
                        # SAFETY: Final cleanup of any embedded newlines
                        clean_content = ' '.join(para.content.split('\n'))
                        clean_content = ' '.join(clean_content.split())
                        run = doc_para.add_run(clean_content)
                        run.font.size = Pt(11)  # Default size

                    logger.debug("Added paragraph with %d characters", len(para.content))

                # Add spacing between paragraphs
                if para_num < len(paragraphs):
                    document.add_paragraph("")
                    logger.debug("Added spacing after paragraph")

            logger.debug("Saving DOCX file to: %s", output_path)
            document.save(output_path)

            file_size = os.path.getsize(output_path) / 1024
            logger.info("DOCX file saved successfully: %s (%.2f KB)",
                       os.path.basename(output_path), file_size)

        except Exception as e:
            logger.error("Error saving paragraphs to DOCX: %s", e, exc_info=True)
            raise
