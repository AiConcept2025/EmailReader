"""
Convert docs to docx
"""
import os
import logging
import re
from typing import Optional
import pdfplumber
from docx import Document

# Optional library for advanced mojibake fixing
try:
    import ftfy
    FTFY_AVAILABLE = True
except ImportError:
    FTFY_AVAILABLE = False

# Get logger for this module
logger = logging.getLogger('EmailReader.DocConverter')

# Regex patterns for deep text sanitization
SURROGATE_PATTERN = re.compile(r'[\ud800-\udfff]')  # UTF-16 surrogates
CONTROL_PATTERN = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]')  # Control chars
C1_CONTROL_PATTERN = re.compile(r'[\u0080-\u009f]')  # C1 control chars
NONCHAR_PATTERN = re.compile(r'[\ufffe\uffff\ufdd0-\ufdef]')  # Unicode non-chars


def sanitize_text_for_xml(text: str) -> str:
    """
    Deep sanitize text content to remove characters that cause
    'Malformed UTF-8 data' errors in vector stores like Flowise.

    This handles:
    - UTF-16 surrogate characters (invalid in UTF-8) - CRITICAL
    - XML 1.0 illegal characters
    - C1 control characters (U+0080-U+009F)
    - Unicode non-characters
    - Mojibake fixing (if ftfy available)

    XML 1.0 allows:
    - #x9 (tab)
    - #xA (line feed)
    - #xD (carriage return)
    - #x20-#xD7FF
    - #xE000-#xFFFD
    - #x10000-#x10FFFF

    Args:
        text: Input text that may contain invalid characters

    Returns:
        Sanitized text safe for UTF-8 encoding and vector stores
    """
    if not text:
        return text

    # Step 1: Try ftfy for mojibake fixing (if available)
    if FTFY_AVAILABLE:
        try:
            text = ftfy.fix_text(text, normalization='NFC')
            logger.debug("Applied ftfy mojibake fixing")
        except Exception as e:
            logger.debug("ftfy fixing failed: %s, continuing with manual fixes", e)

    # Step 2: Remove UTF-16 surrogates (CRITICAL - main cause of error)
    original_length = len(text)
    text = SURROGATE_PATTERN.sub('', text)
    if len(text) != original_length:
        logger.debug("Removed %d UTF-16 surrogate characters", original_length - len(text))

    # Step 3: Remove XML illegal control characters
    original_length = len(text)
    text = CONTROL_PATTERN.sub('', text)
    if len(text) != original_length:
        logger.debug("Removed %d control characters", original_length - len(text))

    # Step 4: Remove C1 control characters
    original_length = len(text)
    text = C1_CONTROL_PATTERN.sub('', text)
    if len(text) != original_length:
        logger.debug("Removed %d C1 control characters", original_length - len(text))

    # Step 5: Remove Unicode non-characters
    original_length = len(text)
    text = NONCHAR_PATTERN.sub('', text)
    if len(text) != original_length:
        logger.debug("Removed %d Unicode non-characters", original_length - len(text))

    # Step 6: UTF-8 round-trip validation
    try:
        # Use surrogatepass to catch any remaining surrogates
        text = text.encode('utf-8', errors='surrogatepass').decode('utf-8', errors='replace')
    except Exception as e:
        logger.debug("UTF-8 surrogatepass failed: %s, using replace", e)
        text = text.encode('utf-8', errors='replace').decode('utf-8', errors='replace')

    # Step 7: Remove replacement characters that may have appeared
    text = text.replace('\ufffd', '')

    # Step 8: Final validation against XML 1.0 spec
    def is_valid_xml_char(char):
        """Check if character is valid in XML 1.0"""
        code_point = ord(char)
        return (
            code_point == 0x09 or  # tab
            code_point == 0x0A or  # line feed
            code_point == 0x0D or  # carriage return
            (0x20 <= code_point <= 0xD7FF) or
            (0xE000 <= code_point <= 0xFFFD) or
            (0x10000 <= code_point <= 0x10FFFF)
        )

    # Filter out any remaining invalid XML characters
    text = ''.join(char for char in text if is_valid_xml_char(char))

    return text


def convert_txt_to_docx(paragraph: str, docx_file_path: str) -> None:
    """
    Converts a plain text file to a Word document. test

    Args:
        paragraph: text.
        docx_file_path: Path to save the output Word document.
    """
    logger.debug("Entering convert_txt_to_docx()")
    logger.debug("Output path: %s", docx_file_path)

    try:
        text_length = len(paragraph)
        logger.debug("Text length: %d characters", text_length)

        if text_length == 0:
            logger.warning("Empty text provided for conversion")

        # Sanitize text to remove XML-incompatible characters
        logger.debug("Sanitizing text for XML compatibility")
        sanitized_paragraph = sanitize_text_for_xml(paragraph)

        sanitized_length = len(sanitized_paragraph)
        if sanitized_length != text_length:
            removed_chars = text_length - sanitized_length
            logger.warning("Removed %d invalid XML characters from text", removed_chars)

        logger.debug("Creating new Word document")
        document = Document()
        document.add_paragraph(sanitized_paragraph)

        logger.debug("Saving document to: %s", docx_file_path)
        document.save(docx_file_path)
        if os.path.exists(docx_file_path):
            file_size = os.path.getsize(docx_file_path) / 1024  # KB
            logger.info('Text converted to Word successfully: %s (%.2f KB)',
                        os.path.basename(docx_file_path), file_size)
        else:
            logger.error(
                "Document save failed - file not found: %s", docx_file_path)

    except Exception as e:
        logger.error("Error converting text to DOCX: %s", e, exc_info=True)
        raise


def convert_txt_file_to_docx(txt_file_path: str, docx_file_path: str) -> None:
    """
    Converts a plain text file to a Word document.

    Args:
        paragraph: text.
        docx_file_path: Path to save the output Word document.
    """
    logger.debug("Entering convert_txt_file_to_docx()")
    logger.debug("Input TXT: %s", txt_file_path)
    logger.debug("Output DOCX: %s", docx_file_path)

    try:
        if not os.path.exists(txt_file_path):
            logger.error("Text file not found: %s", txt_file_path)
            raise FileNotFoundError(f"File not found: {txt_file_path}")

        input_size = os.path.getsize(txt_file_path) / 1024  # KB
        logger.debug("Input file size: %.2f KB", input_size)

        logger.debug("Reading text file")
        with open(txt_file_path, 'r', encoding='utf-8') as file:
            paragraph = file.read()

        text_length = len(paragraph)
        logger.debug("Read %d characters from file", text_length)

        # Sanitize text to remove XML-incompatible characters
        logger.debug("Sanitizing text for XML compatibility")
        sanitized_paragraph = sanitize_text_for_xml(paragraph)

        sanitized_length = len(sanitized_paragraph)
        if sanitized_length != text_length:
            removed_chars = text_length - sanitized_length
            logger.warning("Removed %d invalid XML characters from text", removed_chars)

        logger.debug("Creating Word document")
        document = Document()
        document.add_paragraph(sanitized_paragraph)

        logger.debug("Saving document to: %s", docx_file_path)
        document.save(docx_file_path)

        if os.path.exists(docx_file_path):
            output_size = os.path.getsize(docx_file_path) / 1024  # KB
            logger.info('TXT file converted to Word: %s (%.2f KB)',
                        os.path.basename(docx_file_path), output_size)
        else:
            logger.error(
                "Document save failed - file not found: %s", docx_file_path)

    except Exception as e:
        logger.error("Error converting TXT file to DOCX: %s", e, exc_info=True)
        raise


def convert_pdf_to_docx(pdf_path: str, docx_path: str):
    """
    Converts a PDF file to a DOCX file, preserving text formatting.
    """
    logger.debug("Entering convert_pdf_to_docx()")
    logger.debug("Input PDF: %s", pdf_path)
    logger.debug("Output DOCX: %s", docx_path)

    try:
        if not os.path.exists(pdf_path):
            logger.error("PDF file not found: %s", pdf_path)
            raise FileNotFoundError(f"File not found: {pdf_path}")

        input_size = os.path.getsize(pdf_path) / 1024  # KB
        logger.debug("Input PDF size: %.2f KB", input_size)

        logger.debug("Opening PDF with pdfplumber")
        document = Document()
        total_text_length = 0

        with pdfplumber.open(pdf_path) as pdf:
            num_pages = len(pdf.pages)
            logger.info("PDF has %d pages", num_pages)

            for page_num, page in enumerate(pdf.pages, 1):
                logger.debug("Processing page %d/%d", page_num, num_pages)
                text = page.extract_text()

                if text:
                    text_length = len(text)
                    total_text_length += text_length
                    logger.debug("Page %d: extracted %d characters", page_num, text_length)

                    # Sanitize text to remove XML-incompatible characters
                    sanitized_text = sanitize_text_for_xml(text)

                    if len(sanitized_text) != text_length:
                        removed = text_length - len(sanitized_text)
                        logger.warning("Page %d: removed %d invalid XML characters", page_num, removed)

                    document.add_paragraph(sanitized_text)
                else:
                    logger.debug("Page %d: no text extracted", page_num)

            logger.info("Total text extracted: %d characters from %d pages",
                        total_text_length, num_pages)

            logger.debug("Saving DOCX file")
            document.save(docx_path)

        if os.path.exists(docx_path):
            output_size = os.path.getsize(docx_path) / 1024  # KB
            logger.info('PDF converted to Word: %s (%.2f KB)',
                        os.path.basename(docx_path), output_size)
        else:
            logger.error(
                "Document save failed - file not found: %s", docx_path)

    except Exception as e:
        logger.error("Error converting PDF to DOCX: %s", e, exc_info=True)
        raise
