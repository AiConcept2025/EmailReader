#!/usr/bin/env python3
"""
Example script demonstrating the translate_paragraphs method.

This shows practical usage of the new batch paragraph translation feature.
"""

import os
import sys
import logging
from typing import List

# Add src directory to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.translation.google_doc_translator import GoogleDocTranslator
from src.utils import load_config

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def example_1_basic_translation():
    """Example 1: Basic paragraph translation."""
    print("\n" + "=" * 80)
    print("EXAMPLE 1: Basic Paragraph Translation")
    print("=" * 80)

    # Load configuration
    config = load_config()
    translation_config = config.get('translation', {}).get('google_doc', {})
    service_account = config.get('google_drive', {}).get('service_account')
    if service_account:
        translation_config['service_account'] = service_account

    # Initialize translator
    translator = GoogleDocTranslator(translation_config)

    # Sample paragraphs (English to Spanish)
    paragraphs = [
        "Welcome to the document translation system.",
        "This system can translate multiple paragraphs efficiently.",
        "Each paragraph is processed in batches to optimize API usage."
    ]

    print("\nOriginal paragraphs (English):")
    for i, para in enumerate(paragraphs, 1):
        print(f"  {i}. {para}")

    # Translate
    translated = translator.translate_paragraphs(
        paragraphs=paragraphs,
        target_lang='es'
    )

    print("\nTranslated paragraphs (Spanish):")
    for i, para in enumerate(translated, 1):
        print(f"  {i}. {para}")


def example_2_with_empty_paragraphs():
    """Example 2: Handling empty paragraphs."""
    print("\n" + "=" * 80)
    print("EXAMPLE 2: Preserving Empty Paragraphs")
    print("=" * 80)

    config = load_config()
    translation_config = config.get('translation', {}).get('google_doc', {})
    service_account = config.get('google_drive', {}).get('service_account')
    if service_account:
        translation_config['service_account'] = service_account

    translator = GoogleDocTranslator(translation_config)

    # Paragraphs with empty lines (common in documents)
    paragraphs = [
        "Section 1: Introduction",
        "",  # Empty paragraph (spacing)
        "This is the introduction section.",
        "It contains important information.",
        "",  # Another empty paragraph
        "Section 2: Details",
        "Here are the details."
    ]

    print("\nOriginal structure:")
    for i, para in enumerate(paragraphs):
        if para:
            print(f"  {i}: {para}")
        else:
            print(f"  {i}: (empty line)")

    # Translate
    translated = translator.translate_paragraphs(
        paragraphs=paragraphs,
        target_lang='fr'
    )

    print("\nTranslated structure (French):")
    for i, para in enumerate(translated):
        if para:
            print(f"  {i}: {para}")
        else:
            print(f"  {i}: (empty line preserved)")

    # Verify empty paragraphs preserved
    assert translated[1] == "", "Empty paragraph 1 not preserved!"
    assert translated[4] == "", "Empty paragraph 2 not preserved!"
    print("\n✓ Empty paragraphs successfully preserved")


def example_3_large_batch():
    """Example 3: Translating large number of paragraphs."""
    print("\n" + "=" * 80)
    print("EXAMPLE 3: Large Batch Translation")
    print("=" * 80)

    config = load_config()
    translation_config = config.get('translation', {}).get('google_doc', {})
    service_account = config.get('google_drive', {}).get('service_account')
    if service_account:
        translation_config['service_account'] = service_account

    translator = GoogleDocTranslator(translation_config)

    # Generate 50 paragraphs
    paragraphs = [
        f"This is paragraph number {i}. It contains sample text for testing."
        for i in range(1, 51)
    ]

    print(f"\nTranslating {len(paragraphs)} paragraphs from English to German")
    print("Using batch size of 15 paragraphs per API call")

    # Translate with custom batch size
    translated = translator.translate_paragraphs(
        paragraphs=paragraphs,
        target_lang='de',
        batch_size=15
    )

    print(f"\n✓ Successfully translated {len(translated)} paragraphs")

    # Show first 3 and last 3
    print("\nFirst 3 results:")
    for i in range(3):
        print(f"  {i+1}. {paragraphs[i]}")
        print(f"     → {translated[i]}")

    print("\nLast 3 results:")
    for i in range(47, 50):
        print(f"  {i+1}. {paragraphs[i]}")
        print(f"     → {translated[i]}")


def example_4_from_docx_file():
    """Example 4: Translate paragraphs from a DOCX file."""
    print("\n" + "=" * 80)
    print("EXAMPLE 4: Translating DOCX Document (Paragraph-by-Paragraph)")
    print("=" * 80)

    try:
        from docx import Document
    except ImportError:
        print("⚠ python-docx not installed. Install with: pip install python-docx")
        print("Skipping this example.")
        return

    config = load_config()
    translation_config = config.get('translation', {}).get('google_doc', {})
    service_account = config.get('google_drive', {}).get('service_account')
    if service_account:
        translation_config['service_account'] = service_account

    translator = GoogleDocTranslator(translation_config)

    # Create a sample document
    print("\nCreating sample document...")
    doc = Document()
    doc.add_heading('Sample Document', 0)
    doc.add_paragraph('This is the first paragraph.')
    doc.add_paragraph('This is the second paragraph with more text.')
    doc.add_paragraph('')  # Empty paragraph
    doc.add_paragraph('This is the third paragraph after a break.')

    sample_path = 'temp_sample.docx'
    doc.save(sample_path)
    print(f"  Saved to: {sample_path}")

    # Extract paragraphs
    print("\nExtracting paragraphs from document...")
    input_doc = Document(sample_path)
    paragraphs = [para.text for para in input_doc.paragraphs]
    print(f"  Found {len(paragraphs)} paragraphs")

    # Translate
    print("\nTranslating to Italian...")
    translated = translator.translate_paragraphs(
        paragraphs=paragraphs,
        target_lang='it'
    )

    # Create output document
    print("\nCreating translated document...")
    output_doc = Document()
    for i, (original_para, translated_text) in enumerate(zip(input_doc.paragraphs, translated)):
        if i == 0:
            # Preserve heading
            output_doc.add_heading(translated_text, 0)
        else:
            new_para = output_doc.add_paragraph(translated_text)
            # Copy style from original
            new_para.style = original_para.style

    output_path = 'temp_sample_translated.docx'
    output_doc.save(output_path)
    print(f"  Saved to: {output_path}")

    print("\n✓ Document translation complete!")
    print(f"  Original: {sample_path}")
    print(f"  Translated: {output_path}")

    # Cleanup
    if os.path.exists(sample_path):
        os.remove(sample_path)
        print(f"\n  Cleaned up: {sample_path}")


def example_5_custom_batch_size():
    """Example 5: Experimenting with different batch sizes."""
    print("\n" + "=" * 80)
    print("EXAMPLE 5: Custom Batch Size Configuration")
    print("=" * 80)

    config = load_config()
    translation_config = config.get('translation', {}).get('google_doc', {})
    service_account = config.get('google_drive', {}).get('service_account')
    if service_account:
        translation_config['service_account'] = service_account

    translator = GoogleDocTranslator(translation_config)

    # 30 paragraphs
    paragraphs = [f"Paragraph {i}" for i in range(1, 31)]

    print(f"\nTotal paragraphs: {len(paragraphs)}")

    # Try different batch sizes
    batch_sizes = [5, 10, 15, 30]

    for batch_size in batch_sizes:
        num_batches = (len(paragraphs) + batch_size - 1) // batch_size
        print(f"\n  Batch size {batch_size}: {num_batches} API calls needed")

    print("\nUsing batch_size=10 for translation...")
    translated = translator.translate_paragraphs(
        paragraphs=paragraphs,
        target_lang='es',
        batch_size=10
    )

    print(f"✓ Translated {len(translated)} paragraphs with batch_size=10")


def main():
    """Run all examples."""
    print("\n" + "#" * 80)
    print("# PARAGRAPH TRANSLATION METHOD - USAGE EXAMPLES")
    print("#" * 80)

    print("\nThese examples demonstrate the new translate_paragraphs() method.")
    print("Note: These examples require valid Google Cloud API credentials.")
    print("      Comment out API-dependent examples if credentials are not available.\n")

    # Run examples that don't require API calls
    try:
        # These would require actual API credentials:
        # example_1_basic_translation()
        # example_2_with_empty_paragraphs()
        # example_3_large_batch()
        # example_4_from_docx_file()
        # example_5_custom_batch_size()

        print("\n" + "=" * 80)
        print("EXAMPLE CODE PROVIDED")
        print("=" * 80)
        print("\nTo run these examples with actual translation:")
        print("1. Ensure you have valid Google Cloud credentials configured")
        print("2. Uncomment the example function calls in main()")
        print("3. Run: python3 example_paragraph_translation.py")
        print("\nAvailable examples:")
        print("  - example_1_basic_translation()")
        print("  - example_2_with_empty_paragraphs()")
        print("  - example_3_large_batch()")
        print("  - example_4_from_docx_file()")
        print("  - example_5_custom_batch_size()")

    except Exception as e:
        logger.error("Error running examples: %s", str(e), exc_info=True)
        return False

    return True


if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
